import os
import docx

paths = [
    r"G:\내 드라이브\[FoxConnect]\[총무]업무\01_부동산_자산관리\01_임대차\01_가산_대륭포스트타워6차\402_403호\부동산_계약관리노트_402_403호.docx",
    r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
]

for p in paths:
    print(f"\n==========================================")
    print(f"Inspecting file at: {p}")
    if not os.path.exists(p):
        print("  [FILE NOT FOUND]")
        continue
    try:
        doc = docx.Document(p)
        print("--- PARAGRAPHS ---")
        for i, paragraph in enumerate(doc.paragraphs):
            runs_str = ", ".join([f"Run('{r.text}', sz={r.font.size.pt if r.font.size else 'None'}, b={r.bold})" for r in paragraph.runs])
            print(f"P{i} [{paragraph.style.name}]: '{paragraph.text}' | {runs_str}")

        print("\n--- TABLES ---")
        for t_idx, table in enumerate(doc.tables):
            print(f"\n--- TABLE {t_idx} (rows={len(table.rows)}, cols={len(table.columns)}) ---")
            for r_idx, row in enumerate(table.rows):
                row_cells = []
                for c_idx, cell in enumerate(row.cells):
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    bg = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') if shd is not None else "None"
                    row_cells.append(f"C{c_idx}(bg={bg}): '{cell.text.strip()}'")
                print(f"  R{r_idx}: " + " | ".join(row_cells))
    except Exception as e:
        print(f"Error inspecting {p}: {e}")
