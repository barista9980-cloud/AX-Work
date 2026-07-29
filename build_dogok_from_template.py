import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

template_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
output_docx_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_강남_도곡로1길23.docx"
output_docx_path_alt = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_강남_도곡로1길23_정의양식.docx"

doc = docx.Document(template_path)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_cell_text_and_style(cell, text, bold=False, font_size=9.0):
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(font_size)
            r.bold = bold

# 1. Update Title (P0)
doc.paragraphs[0].text = "🏢 [강남_도곡로1길23] 부동산 계약 관리 노트"
for r in doc.paragraphs[0].runs:
    r.font.name = "맑은 고딕"
    r.font.size = Pt(15.0)
    r.bold = True

# 2. Update Sub-note (P1)
doc.paragraphs[1].text = "※ [AI 파싱 완료] 계약서 PDF에서 OCR 정밀 추출된 내용 입니다.\n※ 최초 작성 시 OCR 정밀 추출 결과물 확인 후 검토해 주시면 됩니다."
for r in doc.paragraphs[1].runs:
    r.font.name = "맑은 고딕"
    r.font.size = Pt(9.0)

# 3. Update Table 0 (주 계약 정보)
t0 = doc.tables[0]

t0_data = [
    [("건물명 / 호수", True), ("강남_도곡로1길23 전층(1~3층 및 지하)", False), ("사용처 (부서/용도)", True), ("강남사무실(폭스에듀 사옥)", False)],
    [("주 계약 유형", True), ("최초임대차", False), ("최초 계약일", True), ("2024-11-07", False)],
    [("임대 기간", True), ("2024-11-30~2026-11-29", False), ("매월 납부일", True), ("", False)],
    [("보증금", True), ("200,000,000", False), ("월 임대료", True), ("", False)],
    [("계약면적 (㎡)", True), ("823.01㎡", False), ("전용면적 / 평수", True), ("823.01㎡/249평", False)]
]

for r_i, row in enumerate(t0_data):
    for c_i, (val, is_header) in enumerate(row):
        cell = t0.rows[r_i].cells[c_i]
        set_cell_text_and_style(cell, val, bold=is_header)
        if is_header:
            set_cell_background(cell, "F2F4F7")

# 4. Update Heading P3 (2. 수록 계약서 문서 목록)
doc.paragraphs[3].text = "2. 수록 계약서 문서 목록 (총 8건)"
for r in doc.paragraphs[3].runs:
    r.font.name = "맑은 고딕"
    r.font.size = Pt(11.0)
    r.bold = True

# 5. Update Table 1 (수록 계약서 문서 목록)
t1 = doc.tables[1]

# Clear existing rows except header
while len(t1.rows) > 1:
    t1._tbl.remove(t1.rows[-1]._tr)

docs_data = [
    ("강남_도곡로1길23_전층_01_최초임대차_[박재윤-㈜폭스에듀]_(241107)", "최초임대차", "박재윤 → ㈜폭스에듀", "2024-11-07"),
    ("강남_도곡로1길23_전층_02_변경계약_[유한회사 청송(박재윤)-㈜폭스에듀]_(250901)", "변경계약", "유한회사 청송(박재윤) → ㈜폭스에듀", "2025-09-01"),
    ("강남_도곡로1길23_1층_01_전대차_[㈜폭스에듀-㈜에스앤에이치트레이딩]_(241101)", "전대차", "㈜폭스에듀 → ㈜에스앤에이치트레이딩", "2024-11-01"),
    ("강남_도곡로1길23_1층_02_전대차_[㈜폭스에듀-한국경찰과학전략센터]_(250821)", "전대차", "㈜폭스에듀 → 한국경찰과학전략센터", "2025-08-21"),
    ("강남_도곡로1길23_1층_03_전대차_[㈜폭스에듀-㈜월드유니코어]_(250821)", "전대차", "㈜폭스에듀 → ㈜월드유니코어", "2025-08-21"),
    ("강남_도곡로1길23_2층_01_전대차_[㈜폭스에듀-㈜실리콘아츠]_(241101)", "전대차", "㈜폭스에듀 → ㈜실리콘아츠", "2024-11-01"),
    ("강남_도곡로1길23_2층_02_전대차_[㈜폭스에듀-㈜하이퍼비주얼에이아이]_(250101)", "전대차", "㈜폭스에듀 → ㈜하이퍼비주얼에이아이", "2025-01-01"),
    ("강남_도곡로1길23_3층_01_전대차_[㈜폭스에듀-㈜트라이디스]_(250124)", "전대차", "㈜폭스에듀 → ㈜트라이디스", "2025-01-24")
]

# Set Header formatting
t1_headers = ["문서명", "계약 종류", "계약 당사자 (임대인 - 임차인)", "계약일"]
for c_i, h in enumerate(t1_headers):
    cell = t1.rows[0].cells[c_i]
    set_cell_text_and_style(cell, h, bold=True)
    set_cell_background(cell, "E6ECF5")

# Add Data Rows
for d_row in docs_data:
    row_cells = t1.add_row().cells
    for c_i, val in enumerate(d_row):
        set_cell_text_and_style(row_cells[c_i], val, bold=False)

# 6. Update Table 2 (임대인 및 납부 계좌 정보)
t2 = doc.tables[2]

t2_data = [
    [("주 임대인", True), ("유한회사 청송(박재윤)", False), ("주 임차인", True), ("㈜폭스에듀(폭스커넥트 법인)", False)],
    [("임대인 연락처", True), ("02-2007-9152", False), ("관리사무소 연락처", True), ("", False)],
    [("입금 은행", True), ("신한은행", False), ("예금주", True), ("유한회사 청송", False)],
    [("계좌번호", True), ("140-015-707061", False), ("비고", True), ("2025-09-01 개인에서 법인(유한회사 청송)으로 변경 승계", False)]
]

for r_i, row in enumerate(t2_data):
    for c_i, (val, is_header) in enumerate(row):
        cell = t2.rows[r_i].cells[c_i]
        set_cell_text_and_style(cell, val, bold=is_header)
        if is_header:
            set_cell_background(cell, "F2F4F7")

# 7. Update Table 3 (계약 변동 이력 및 특이사항)
t3 = doc.tables[3]

history_items = [
    "2025-09-01[변경계약] 박재윤 → 유한회사 청송(임대인 승계)",
    "2024-11-01[전대차] ㈜폭스에듀 → ㈜에스앤에이치트레이딩(1층 전체)",
    "2025-08-21[전대차] ㈜폭스에듀 → 한국경찰과학전략센터(1층 일부)",
    "2025-08-21[전대차] ㈜폭스에듀 → ㈜월드유니코어(1층 일부)",
    "2024-11-01[전대차] ㈜폭스에듀 → ㈜실리콘아츠(2층 전체)",
    "2025-01-01[전대차] ㈜폭스에듀 → ㈜하이퍼비주얼에이아이(2층 전체)",
    "2025-01-24[전대차] ㈜폭스에듀 → ㈜트라이디스(3층 전체)"
]
history_text = "\n".join(history_items)

t3_data = [
    [("계약 변동 / 전대차 / 승계 이력", True), (history_text, False)],
    [("계약 연장 / 묵시적 갱신 이력", True), ("", False)],
    [("중도해지 / 퇴거 예정 메모", True), ("", False)],
    [("기타 특약 및 참조사항", True), ("", False)]
]

for r_i, row in enumerate(t3_data):
    for c_i, (val, is_header) in enumerate(row):
        cell = t3.rows[r_i].cells[c_i]
        set_cell_text_and_style(cell, val, bold=is_header)
        if is_header:
            set_cell_background(cell, "F2F4F7")

saved_path = ""
try:
    doc.save(output_docx_path)
    saved_path = output_docx_path
except Exception as e:
    doc.save(output_docx_path_alt)
    saved_path = output_docx_path_alt

print("Saved perfectly matched docx to:", saved_path)
