import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
GANGSAN_DIR = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리\01_임대차계약\22_강남_도곡로1길23_지하1층,1층,2층,3층")
docx_p = os.path.join(GANGSAN_DIR, "부동산_계약관리노트_강남_도곡로1길23_지하1층,1층,2층,3층.docx")

if os.path.exists(docx_p):
    print("Updating Gangsan Word Contract Note with Verified Figures...")
    doc = docx.Document(docx_p)
    
    # Table 1 update
    if len(doc.tables) > 0:
        t1 = doc.tables[0]
        # find deposit cell
        for row in t1.rows:
            for i, cell in enumerate(row.cells):
                if "보증금" in cell.text and i + 1 < len(row.cells):
                    row.cells[i+1].text = "200,000,000원 (2억원)"
                if "월 임대료" in cell.text and i + 1 < len(row.cells):
                    row.cells[i+1].text = "14,700,000원 (월 1,470만원, VAT별도 / 관리비 2,190,800원)"

    doc.save(docx_p)
    print("  [WORD CONTRACT NOTE UPDATED SUCCESSFULLY]")
