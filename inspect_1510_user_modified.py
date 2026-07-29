import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

file_1510_path = None
for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    if "1510" in root:
        for f in files:
            if f.endswith(".docx") and "계약관리노트" in f:
                file_1510_path = os.path.join(root, f)
                break
    if file_1510_path:
        break

print("Found 1510호 docx file at:", file_1510_path)

if file_1510_path:
    doc = docx.Document(file_1510_path)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P[{idx}]: {p.text}")
            
    print(f"\nTotal Tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} (Rows={len(table.rows)}, Cols={len(table.columns)}) ---")
        for r_idx, row in enumerate(table.rows):
            cell_texts = []
            for c_idx, cell in enumerate(row.cells):
                cell_texts.append(f"C{c_idx}: '{cell.text.replace('\n', ' // ')}'")
            print(f"  Row {r_idx}: " + " | ".join(cell_texts))
