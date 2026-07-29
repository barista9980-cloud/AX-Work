import docx

doc_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
doc = docx.Document(doc_path)

print("=== PARAGRAPHS IN 402_403호 ===")
for i, p in enumerate(doc.paragraphs):
    print(f"P{i}: '{p.text}'")

for t_idx, t in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, r in enumerate(t.rows):
        for c_idx, c in enumerate(r.cells):
            if "." in c.text or "•" in c.text or "·" in c.text:
                print(f"  T{t_idx} R{r_idx} C{c_idx}: '{c.text}'")
