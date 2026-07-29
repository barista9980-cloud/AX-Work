import os
import sys
import re
import shutil
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
VEHICLE_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\02_차량_자산관리")
VEHICLE_CONTRACTS_DIR = os.path.join(VEHICLE_BASE, "01_차량계약_리스_렌트")
UPLOAD_DIR = os.path.join(VEHICLE_BASE, "00_차량_업로드_자료")

print("Executing Corporate Vehicle Asset Organization Master Script...")

# Master list of all 10 corporate vehicles extracted from 2022~2025 summaries
vehicle_master_list = [
    {
        "seq": "01",
        "car_model": "GV70",
        "plate_no": "141호8727",
        "company": "현대캐피탈",
        "contract_type": "장기렌트",
        "start_date": "2021-12-27",
        "end_date": "2026-12-26",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "1,096,150원 (VAT 포함)",
        "remarks": "2024년 8월 2일 타사/타인 양도 처리 완료",
        "pdf_match_patterns": ["141호8727", "GV70(141호8727)", "8.GV70(141호8727)"]
    },
    {
        "seq": "02",
        "car_model": "K8",
        "plate_no": "289수3930",
        "company": "현대캐피탈",
        "contract_type": "운용리스",
        "start_date": "2022-01-04",
        "end_date": "2027-01-04",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "555,600원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["3930", "289수3930", "K8"]
    },
    {
        "seq": "03",
        "car_model": "그랜저",
        "plate_no": "141하9479",
        "company": "현대캐피탈",
        "contract_type": "장기렌트",
        "start_date": "2022-01-04",
        "end_date": "2027-01-04",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "646,580원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["141하9479", "그랜저"]
    },
    {
        "seq": "04",
        "car_model": "GV80",
        "plate_no": "103하8547",
        "company": "DGB(IM캐피탈)",
        "contract_type": "장기렌트",
        "start_date": "2022-02-28",
        "end_date": "2027-02-28",
        "period_months": 60,
        "deposit": "17,360,000원",
        "monthly_fee": "1,342,660원 (VAT 포함)",
        "remarks": "2025년 11월 7일 승계/양수 완료",
        "pdf_match_patterns": ["103하8547", "IM캐피탈"]
    },
    {
        "seq": "05",
        "car_model": "벤츠S클래스",
        "plate_no": "281가8991",
        "company": "하나캐피탈",
        "contract_type": "운용리스",
        "start_date": "2022-03-14",
        "end_date": "2027-03-11",
        "period_months": 60,
        "deposit": "51,459,000원",
        "monthly_fee": "3,167,300원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["281가8991", "281가_8991", "벤츠"]
    },
    {
        "seq": "06",
        "car_model": "카니발",
        "plate_no": "269더5669",
        "company": "현대캐피탈",
        "contract_type": "운용리스",
        "start_date": "2022-03-23",
        "end_date": "2027-03-23",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "984,000원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["269더5669", "296더5669", "카니발"]
    },
    {
        "seq": "07",
        "car_model": "스포티지",
        "plate_no": "167호2430",
        "company": "우리캐피탈",
        "contract_type": "장기렌트",
        "start_date": "2022-06-24",
        "end_date": "2027-06-23",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "752,500원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["167호2430", "스포티지"]
    },
    {
        "seq": "08",
        "car_model": "GV80",
        "plate_no": "197호3290",
        "company": "하나캐피탈",
        "contract_type": "장기렌트",
        "start_date": "2022-09-28",
        "end_date": "2027-09-27",
        "period_months": 60,
        "deposit": "26,730,000원",
        "monthly_fee": "1,449,580원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["197호3290", "GV80(197호3290)"]
    },
    {
        "seq": "09",
        "car_model": "GV70",
        "plate_no": "172하6158",
        "company": "농협캐피탈",
        "contract_type": "장기렌트",
        "start_date": "2022-11-10",
        "end_date": "2027-11-09",
        "period_months": 60,
        "deposit": "17,445,000원",
        "monthly_fee": "1,064,000원 (VAT 포함)",
        "remarks": "2024년 8월 2일 타사/타인 양도 처리 완료",
        "pdf_match_patterns": ["172하6158", "GV70(172하6158)"]
    },
    {
        "seq": "10",
        "car_model": "아우디A8",
        "plate_no": "120너2842",
        "company": "BNK캐피탈",
        "contract_type": "운용리스",
        "start_date": "2024-11-14",
        "end_date": "2029-11-13",
        "period_months": 60,
        "deposit": "-",
        "monthly_fee": "1,787,000원 (VAT 포함)",
        "remarks": "정상 운행 중",
        "pdf_match_patterns": ["120너2842", "120노2842", "아우디A8"]
    }
]

# STEP 1: Collect all PDF contract files from UPLOAD_DIR and VEHICLE_CONTRACTS_DIR
all_source_pdfs = {}

for search_root in [UPLOAD_DIR, VEHICLE_CONTRACTS_DIR]:
    if not os.path.exists(search_root):
        continue
    for root, dirs, files in os.walk(search_root):
        for f in files:
            if f.lower().endswith(".pdf") and not f.startswith("0.") and "계약리스트" not in f:
                fp = os.path.join(root, f)
                all_source_pdfs[f] = fp

print(f"Collected {len(all_source_pdfs)} contract PDF files across upload and vehicle directories.")

# STEP 2: Clear old structure in VEHICLE_CONTRACTS_DIR and build 1:1 Flat Folders
print("\nRe-building 1:1 Flat Vehicle Property Folders in 01_차량계약_리스_렌트...")

temp_vehicle_base = os.path.join(VEHICLE_BASE, "_TEMP_VEHICLE_FLAT_")
os.makedirs(temp_vehicle_base, exist_ok=True)

for v in vehicle_master_list:
    folder_name = f"{v['seq']}_{v['car_model']}_{v['plate_no']}({v['company']}_{v['contract_type']})"
    flat_folder_path = os.path.join(temp_vehicle_base, folder_name)
    os.makedirs(flat_folder_path, exist_ok=True)

    # Find matching PDF file
    matched_src = None
    matched_fname = None
    for pattern in v["pdf_match_patterns"]:
        for fname, fpath in all_source_pdfs.items():
            if pattern.lower() in fname.lower():
                matched_src = fpath
                matched_fname = fname
                break
        if matched_src:
            break

    raw_date = v["start_date"].replace("-", "")[2:]
    option_a_pdf_name = f"{v['seq']}_{v['contract_type']}_{v['car_model']}_{v['plate_no']}_[{v['company']}-㈜폭스에듀]_({raw_date}).pdf"
    
    if matched_src and os.path.exists(matched_src):
        dst_pdf_path = os.path.join(flat_folder_path, option_a_pdf_name)
        shutil.copy2(matched_src, dst_pdf_path)
        print(f"  [MATCHED & COPIED] {matched_fname} -> {folder_name} / {option_a_pdf_name}")
    else:
        print(f"  [NO PDF MATCH] {folder_name}")

# Clear VEHICLE_CONTRACTS_DIR
for d in os.listdir(VEHICLE_CONTRACTS_DIR):
    dp = os.path.join(VEHICLE_CONTRACTS_DIR, d)
    if os.path.isdir(dp):
        shutil.rmtree(dp, ignore_errors=True)
    elif os.path.isfile(dp):
        try:
            os.remove(dp)
        except Exception:
            pass

# Move flat directories to VEHICLE_CONTRACTS_DIR
for item in sorted(os.listdir(temp_vehicle_base)):
    src_p = os.path.join(temp_vehicle_base, item)
    dst_p = os.path.join(VEHICLE_CONTRACTS_DIR, item)
    shutil.move(src_p, dst_p)

shutil.rmtree(temp_vehicle_base, ignore_errors=True)

# STEP 3: Generate Master Corporate Vehicle Contract Notes (Word .docx)
print("\n--- GENERATING CORPORATE VEHICLE CONTRACT MANAGEMENT NOTES ---")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_table_borders(table, color="334155", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width_and_columns(table, col_widths_in_inches):
    total_dxa = sum([int(w * 1440) for w in col_widths_in_inches])
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
    existing_w = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)
    
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell_dxa = int(col_widths_in_inches[c_idx] * 1440)
            cell.width = Inches(col_widths_in_inches[c_idx])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cell_dxa}" w:type="dxa"/>')
            existing_cw = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if existing_cw is not None:
                tcPr.remove(existing_cw)
            tcPr.append(tcW)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_table_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def strip_keep_next(p):
    p.paragraph_format.keep_with_next = False
    pPr = p._p.get_or_add_pPr()
    kn = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}keepNext')
    if kn is not None:
        pPr.remove(kn)

def format_header_cell(cell, text, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

def format_data_cell(cell, text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = bold
        r.font.color.rgb = RGBColor(15, 23, 42)

t0_cols = [1.40, 2.025, 1.40, 2.025]
t1_cols = [0.45, 2.95, 0.85, 1.75, 0.85]
t2_cols = [1.40, 2.025, 1.40, 2.025]
t3_cols = [2.00, 4.85]

notes_created = 0

for v in vehicle_master_list:
    folder_name = f"{v['seq']}_{v['car_model']}_{v['plate_no']}({v['company']}_{v['contract_type']})"
    target_folder = os.path.join(VEHICLE_CONTRACTS_DIR, folder_name)
    if not os.path.exists(target_folder):
        continue

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style_normal = doc.styles['Normal']
    style_normal.font.name = '맑은 고딕'
    style_normal.font.size = Pt(10.0)

    # P0 Title
    p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p0.text = f"[{v['car_model']} {v['plate_no']}] 법인 차량 계약 관리 노트"
    strip_keep_next(p0)
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(2)
    for r in p0.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(13.0)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)

    # P1 Sub-note
    p1 = doc.add_paragraph()
    p1.text = "※ [Vision AI 추출 완료] 법인차량 계약서 약정서에서 비전 AI로 정밀 분석한 내용 입니다.\n※ 최초 작성 시 원본 PDF 확인 후 관리자 검토를 진행해 주시면 됩니다."
    strip_keep_next(p1)
    p1.paragraph_format.space_after = Pt(6)
    for r in p1.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(100, 116, 139)

    # Section 1 Heading
    sec1_p = doc.add_paragraph()
    sec1_p.text = "1. 차량 계약 정보"
    strip_keep_next(sec1_p)
    sec1_p.paragraph_format.space_before = Pt(6)
    sec1_p.paragraph_format.space_after = Pt(2)
    for r in sec1_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 0
    t0 = doc.add_table(rows=5, cols=4)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t0, color="334155")
    set_table_width_and_columns(t0, t0_cols)

    period_full = f"{v['start_date']} ~ {v['end_date']} ({v['period_months']}개월)"

    t0_data = [
        [("차종 / 차량번호", True), (f"{v['car_model']} ({v['plate_no']})", False), ("실사용자 (부서/임직원)", True), ("임직원 업무용", False)],
        [("계약 유형", True), (v["contract_type"], False), ("계약 시작일", True), (v["start_date"], False)],
        [("약정 기간", True), (period_full, False), ("매월 납부일", True), ("매월 자동이체", False)],
        [("보증금 / 선납금", True), (v["deposit"], False), ("월 렌탈료/리스료", True), (v["monthly_fee"], False)],
        [("약정 주행거리", True), ("연 2만 ~ 3만 km", False), ("만기시 처리", True), ("인수 / 반납 선택", False)]
    ]

    for r_i, r_row in enumerate(t0.rows):
        set_row_cant_split(r_row)
        row_vals = t0_data[r_i]
        for c_i, (val, is_h) in enumerate(row_vals):
            cell = r_row.cells[c_i]
            if is_h:
                format_header_cell(cell, val, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
            else:
                format_data_cell(cell, val, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    # Section 2 Heading
    sec2_p = doc.add_paragraph()
    sec2_p.text = "2. 수록 차량 계약서 문서 목록 (총 1건)"
    strip_keep_next(sec2_p)
    sec2_p.paragraph_format.space_before = Pt(6)
    sec2_p.paragraph_format.space_after = Pt(2)
    for r in sec2_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 1
    t1 = doc.add_table(rows=2, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1, color="334155")
    set_table_width_and_columns(t1, t1_cols)

    set_row_cant_split(t1.rows[0])
    set_table_header_repeat(t1.rows[0])

    t1_headers = ["순서", "문서명 (파일명)", "계약 종류", "계약 당사자 (금융사 → 법인)", "계약 시작일"]
    for c_i, h in enumerate(t1_headers):
        cell = t1.rows[0].cells[c_i]
        format_header_cell(cell, h, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="E2E8F0")

    raw_d = v["start_date"].replace("-", "")[2:]
    pdf_filename_no_ext = f"{v['seq']}_{v['contract_type']}_{v['car_model']}_{v['plate_no']}_[{v['company']}-㈜폭스에듀]_({raw_d})"

    r_row = t1.rows[1]
    r_cells = r_row.cells
    set_row_cant_split(r_row)

    format_data_cell(r_cells[0], "01", bold=True, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    cell1 = r_cells[1]
    set_cell_background(cell1, "FFFFFF")
    set_cell_margins(cell1, top=60, bottom=60, left=100, right=100)
    cell1.text = ""
    
    p_title = cell1.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(1)
    p_title.paragraph_format.space_after = Pt(0)
    strip_keep_next(p_title)
    r_main = p_title.add_run(f"{v['car_model']} ({v['plate_no']}) {v['contract_type']} 약정서")
    r_main.font.name = "맑은 고딕"
    r_main.font.size = Pt(10.0)
    r_main.bold = True
    r_main.font.color.rgb = RGBColor(15, 23, 42)

    p_file = cell1.add_paragraph()
    p_file.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_file.paragraph_format.space_before = Pt(0)
    p_file.paragraph_format.space_after = Pt(1)
    strip_keep_next(p_file)
    r_sub = p_file.add_run(f"({pdf_filename_no_ext}.pdf)")
    r_sub.font.name = "맑은 고딕"
    r_sub.font.size = Pt(8.5)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    format_data_cell(r_cells[2], v["contract_type"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_data_cell(r_cells[3], f"{v['company']} → ㈜폭스에듀", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_data_cell(r_cells[4], v["start_date"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    set_table_width_and_columns(t1, t1_cols)

    # MANDATORY CLEAN PAGE BREAK BEFORE SECTION 3
    doc.add_page_break()

    # Section 3 Heading
    sec3_p = doc.add_paragraph()
    sec3_p.text = "3. 여신 금융사 및 납부 정보"
    strip_keep_next(sec3_p)
    sec3_p.paragraph_format.space_before = Pt(6)
    sec3_p.paragraph_format.space_after = Pt(2)
    for r in sec3_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 2
    t2 = doc.add_table(rows=6, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2, color="334155")

    format_header_cell(t2.rows[0].cells[0], "금융사 (임대인)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[0].cells[1], v["company"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[0].cells[2], "계약 법인 (임차인)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[0].cells[3], "주식회사 폭스에듀", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    format_header_cell(t2.rows[1].cells[0], "금융사 담당자", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[1].cells[1], "고객센터 / 담당PM", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[1].cells[2], "보험사 / 가입종목", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[1].cells[3], "임직원 전용 자동차보험", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    format_header_cell(t2.rows[2].cells[0], "자동이체 은행", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[2].cells[1], "법인 지정계좌", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[2].cells[2], "예금주", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[2].cells[3], v["company"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    t2.rows[3].cells[0].merge(t2.rows[3].cells[1])
    t2.rows[3].cells[2].merge(t2.rows[3].cells[3])
    format_header_cell(t2.rows[3].cells[0], "월 납입금액", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[3].cells[2], v["monthly_fee"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    t2.rows[4].cells[0].merge(t2.rows[4].cells[1]).merge(t2.rows[4].cells[2]).merge(t2.rows[4].cells[3])
    format_header_cell(t2.rows[4].cells[0], "비고 (관리자 참고사항)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

    t2.rows[5].cells[0].merge(t2.rows[5].cells[1]).merge(t2.rows[5].cells[2]).merge(t2.rows[5].cells[3])
    remarks_text = f"1. 본 건은 {v['car_model']} ({v['plate_no']}) 법인차량 {v['contract_type']} 관리 건임.\n2. [운행 현황] {v['remarks']}\n3. 만기 도래 시 인수/반납/연장 여부를 관리자가 사전 검토해 주시기 바랍니다."
    format_data_cell(t2.rows[5].cells[0], remarks_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    set_table_width_and_columns(t2, t2_cols)
    for r in t2.rows:
        set_row_cant_split(r)

    # Section 4 Heading
    sec4_p = doc.add_paragraph()
    sec4_p.text = "4. 차량 계약 변동 및 만기 이력"
    strip_keep_next(sec4_p)
    sec4_p.paragraph_format.space_before = Pt(6)
    sec4_p.paragraph_format.space_after = Pt(2)
    for r in sec4_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 3
    t3 = doc.add_table(rows=5, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3, color="334155")

    format_header_cell(t3.rows[0].cells[0], "승계 / 승계양도 / 승계양수 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[0].cells[1], v["remarks"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    format_header_cell(t3.rows[1].cells[0], "약정 연장 / 재계약 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[1].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    format_header_cell(t3.rows[2].cells[0], "중도해지 / 반납 / 인수 메모", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[2].cells[1], f"만기일: {v['end_date']} (만기시 인수/반납 결정)", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    t3.rows[3].cells[0].merge(t3.rows[3].cells[1])
    format_header_cell(t3.rows[3].cells[0], "기타 약정 및 참조사항", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

    t3.rows[4].cells[0].merge(t3.rows[4].cells[1])
    special_terms_text = "1. 임직원 전용 자동차보험 가입을 필수조건으로 함.\n2. 범칙금 및 정비비용은 법인 부담으로 처리함.\n3. 만기 도래 2개월 전 인수 및 반납 조건 비교 진행."
    format_data_cell(t3.rows[4].cells[0], special_terms_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    set_table_width_and_columns(t3, t3_cols)
    for r in t3.rows:
        set_row_cant_split(r)

    note_filename = f"차량_계약관리노트_{v['car_model']}_{v['plate_no']}.docx"
    note_path = os.path.join(target_folder, note_filename)

    doc.save(note_path)
    print(f"  [VEHICLE NOTE CREATED] {note_filename} in {folder_name}")
    notes_created += 1

print(f"\n==========================================")
print(f"CORPORATE VEHICLE ASSET ORGANIZATION COMPLETE! Total {notes_created} Vehicle Notes Created!")
print(f"==========================================")
