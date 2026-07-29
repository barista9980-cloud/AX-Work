import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

doc = docx.Document()

title = doc.add_heading("🏢 [강남_도곡로1길23] 부동산 계약 관리 노트", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("※ 본 문서는 [FoxConnect] 구글 드라이브 02_강남_도곡로1길23 폴더의 원천 계약서 8건을 정밀 분석하여 작성된 통합 계약 관리 노트입니다.")
p.runs[0].font.size = Pt(9.5)
p.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading("1. 건물의 기본 마스터 데이터 (Master Data)", level=2)
t1 = doc.add_table(rows=7, cols=2)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER

m_data = [
    ("건물명 / 소재지", "강남_도곡로1길23 (서울특별시 강남구 도곡로1길 23)"),
    ("건물 연면적", "823.01 ㎡ (지하 1층, 지상 1~3층)"),
    ("주 임대인 (소유주)", "유한회사 청송 (대표자: 박재윤) ※ 2025-09-01 법인 승계"),
    ("주 임차인", "㈜폭스에듀"),
    ("최초 계약일 / 변경일", "2024-11-07 (효력: 2024.11.30~) / 2025-09-01"),
    ("주 보증금", "200,000,000 원 (2억 원)"),
    ("주 임대료 입금 계좌", "신한은행 140-015-707061 (예금주: 유한회사 청송)")
]

for idx, (k, v) in enumerate(m_data):
    row = t1.rows[idx]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_background(row.cells[0], "F2F4F7")

doc.add_heading("2. 수록 계약서 문서 목록 (총 8건)", level=2)
t2 = doc.add_table(rows=9, cols=5)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["순서", "구분", "계약 종류", "계약 당사자 (임대/전대 → 임차/전차)", "계약일"]
for c_i, h in enumerate(headers):
    t2.rows[0].cells[c_i].text = h
    set_cell_background(t2.rows[0].cells[c_i], "E6ECF5")

docs_list = [
    ("01", "전층", "최초임대차", "박재윤 → ㈜폭스에듀", "2024-11-07"),
    ("02", "전층", "변경계약", "유한회사 청송(박재윤) → ㈜폭스에듀", "2025-09-01"),
    ("03", "1층", "전대차 (전체)", "㈜폭스에듀 → ㈜에스앤에이치트레이딩", "2024-11-01"),
    ("04", "1층", "전대차 (일부)", "㈜폭스에듀 → 한국경찰과학전략센터", "2025-08-21"),
    ("05", "1층", "전대차 (일부)", "㈜폭스에듀 → ㈜월드유니코어", "2025-08-21"),
    ("06", "2층", "전대차", "㈜폭스에듀 → ㈜실리콘아츠", "2024-11-01"),
    ("07", "2층", "전대차", "㈜폭스에듀 → ㈜하이퍼비주얼에이아이", "2025-01-01"),
    ("08", "3층", "전대차", "㈜폭스에듀 → ㈜트라이디스", "2025-01-24")
]

for r_i, r_data in enumerate(docs_list, 1):
    row = t2.rows[r_i]
    for c_i, val in enumerate(r_data):
        row.cells[c_i].text = val

doc.add_heading("3. 층별 세부 계약 및 변경 이력", level=2)
p_hist = doc.add_paragraph(
    "[메인 임대차] 전층 계약\n"
    "• 최초 계약 (문서 01): 2024-11-07, 임대인 박재윤 개인과 ㈜폭스에듀 간 전층(823.01㎡) 최초 임대차 계약 체결. 보증금 2억 원. (입금계좌: 신한은행 140-008-769560)\n"
    "• 임대인 변경 승계 (문서 02): 2025-09-01, 건물 소유주가 '유한회사 청송'으로 변경됨에 따라 임대인 승계 변경계약 체결. (변경 입금계좌: 신한은행 140-015-707061)\n\n"
    "[1층 전대차 이력]\n"
    "• 문서 03: 2024-11-01 ㈜에스앤에이치트레이딩 (1층 전체 191.01㎡, 입금계좌: 농협 356-118457-04-015)\n"
    "• 문서 04: 2025-08-21 한국경찰과학전략센터 (1층 일부 95.50㎡)\n"
    "• 문서 05: 2025-08-21 ㈜월드유니코어 (1층 일부 95.50㎡)\n\n"
    "[2층 전대차 이력]\n"
    "• 문서 06: 2024-11-01 ㈜실리콘아츠 (2층 전체 191.01㎡)\n"
    "• 문서 07: 2025-01-01 ㈜하이퍼비주얼에이아이 (2층 전체 191.01㎡)\n\n"
    "[3층 전대차 이력]\n"
    "• 문서 08: 2025-01-24 ㈜트라이디스 (3층 전대차)"
)

doc.add_heading("4. 🚨 데이터 판독 및 분석 결과 보고 (Anti-Hallucination)", level=2)
p_err = doc.add_paragraph(
    "🚨 오류/누락: 전대차 계약서 일부 스캔본(문서 03~08)의 특약사항 및 별첨 페이지의 스캔 해상도가 낮아 각 전차인별 정확한 월차임 및 임대만료일 단독 텍스트가 원본 이미지상에서 손상되어 있습니다.\n"
    "🛠️ 요청 사항: 1층~3층 전차인과의 전대차 계약서 최종 특약 페이지 또는 전대차 입금/세금계산서 내역을 첨부해 주시면, 마스터 데이터 표에 100% 정밀 반영이 가능합니다."
)

target_docx_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_강남_도곡로1길23.docx"
doc.save(target_docx_path)
print("Saved docx to:", target_docx_path)
