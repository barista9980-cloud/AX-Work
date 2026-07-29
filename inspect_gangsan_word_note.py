import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
GANGSAN_DIR = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리\01_임대차계약\22_강남_도곡로1길23_지하1층,1층,2층,3층")

docx_p = os.path.join(GANGSAN_DIR, "부동산_계약관리노트_강남_도곡로1길23.docx")

if os.path.exists(docx_p):
    print("Reading Word Contract Note:", docx_p)
    doc = docx.Document(docx_p)
    for p in doc.paragraphs:
        if p.text.strip():
            print("P:", p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            row_txt = [c.text.strip().replace("\n", " ") for c in r.cells]
            print("TBL:", " | ".join(row_txt))
else:
    print("Word Note NOT FOUND at:", docx_p)
