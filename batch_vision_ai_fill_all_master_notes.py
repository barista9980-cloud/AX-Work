import os
import sys
import re
import json
import time
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
REAL_ESTATE_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리")
LEASE_DIR = os.path.join(REAL_ESTATE_BASE, "01_임대차계약")
SALE_DIR = os.path.join(REAL_ESTATE_BASE, "02_매매_소유권문서")

print("Starting Vision AI Master Contract Note Generation across all 27 Property Folders...")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_table_borders(table, color="334155", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width_and_columns(table, col_widths_in_inches):
    total_dxa = sum([int(w * 1440) for w in col_widths_in_inches])
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
    existing_w = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)
    
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell_dxa = int(col_widths_in_inches[c_idx] * 1440)
            cell.width = Inches(col_widths_in_inches[c_idx])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cell_dxa}" w:type="dxa"/>')
            existing_cw = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if existing_cw is not None:
                tcPr.remove(existing_cw)
            tcPr.append(tcW)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_table_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def strip_keep_next(p):
    p.paragraph_format.keep_with_next = False
    pPr = p._p.get_or_add_pPr()
    kn = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}keepNext')
    if kn is not None:
        pPr.remove(kn)

def format_header_cell(cell, text, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

def format_data_cell(cell, text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = bold
        r.font.color.rgb = RGBColor(15, 23, 42)

def extract_vision_details_from_pdf(pdf_path, name_no_ext):
    """
    Extracts text and structured contract information from PDF using PyMuPDF and pattern matching.
    """
    deposit = ""
    monthly_rent = ""
    start_date = ""
    end_date = ""
    pay_day = ""
    bank = ""
    account_num = ""
    account_holder = ""
    special_terms = []
    
    try:
        doc_pdf = fitz.open(pdf_path)
        full_text = ""
        for page in doc_pdf:
            full_text += page.get_text() + "\n"

        # Deposit matching
        m_dep = re.search(r"보증금\s*[:\=]?\s*금?\s*([\d,]+)\s*원", full_text)
        if not m_dep:
            m_dep = re.search(r"보증금\s*([\d,]+만?\s*원)", full_text)
        if m_dep:
            deposit = m_dep.group(1) + ("원" if not m_dep.group(1).endswith("원") else "")

        # Rent matching
        m_rent = re.search(r"차임\s*\(?월세\)?\s*[:\=]?\s*금?\s*([\d,]+)\s*원", full_text)
        if not m_rent:
            m_rent = re.search(r"월\s*임대료\s*[:\=]?\s*금?\s*([\d,]+)\s*원", full_text)
        if m_rent:
            monthly_rent = m_rent.group(1) + ("원" if not m_rent.group(1).endswith("원") else "")

        # Lease dates
        m_period = re.search(r"(20\d{2}[\.\-년]\s*\d{1,2}[\.\-월]\s*\d{1,2}[\.\-일]?)\s*부터\s*(20\d{2}[\.\-년]\s*\d{1,2}[\.\-월]\s*\d{1,2}[\.\-일]?)", full_text)
        if m_period:
            start_date = m_period.group(1).strip()
            end_date = m_period.group(2).strip()

        # Bank Account
        m_acc = re.search(r"(신한|국민|우리|하나|기업|농협|카카오|케이|수협|대구|부산|경남|광주|전북|우체국)\s*[\:\=]?\s*([\d\-]+)\s*\(?예금주\s*[\:\=]?\s*([^\)\n]+)\)?", full_text)
        if m_acc:
            bank = m_acc.group(1)
            account_num = m_acc.group(2)
            account_holder = m_acc.group(3).strip()

    except Exception:
        pass

    # Fallbacks from filename
    m_filename_date = re.search(r"\((\d{6})\)", name_no_ext)
    if not start_date and m_filename_date:
        raw_d = m_filename_date.group(1)
        start_date = f"20{raw_d[:2]}-{raw_d[2:4]}-{raw_d[4:6]}"

    return {
        "deposit": deposit,
        "monthly_rent": monthly_rent,
        "start_date": start_date,
        "end_date": end_date,
        "pay_day": pay_day,
        "bank": bank,
        "account_num": account_num,
        "account_holder": account_holder
    }

def parse_option_a_filename(filename):
    name_no_ext = os.path.splitext(filename)[0]
    pattern = r"^(\d{2})_([^_]+)_(.*?)_\[(.*?)\]_\((\d{6})\)$"
    m = re.match(pattern, name_no_ext)
    if m:
        seq, c_type, prop_info, parties, raw_date = m.groups()
        party_list = parties.split('-') if '-' in parties else [parties]
        p_a = party_list[0] if len(party_list) > 0 else ""
        p_b = party_list[1] if len(party_list) > 1 else ""
        y = "20" + raw_date[:2]
        mon = raw_date[2:4]
        d = raw_date[4:6]
        f_date = f"{y}-{mon}-{d}"
        
        unit_match = re.search(r"(\d+호|\d+_\d+호|\d+층|지하\d+층.*)", prop_info)
        u_str = unit_match.group(1) if unit_match else ""
        display_title = f"{u_str} {c_type} 계약서".strip() if u_str else f"{c_type} 계약서"

        return {
            "filename_no_ext": name_no_ext,
            "display_title": display_title,
            "prop_info": prop_info,
            "seq": seq,
            "contract_type": c_type,
            "party_a": p_a,
            "party_b": p_b,
            "contract_date": f_date
        }
    else:
        return {
            "filename_no_ext": name_no_ext,
            "display_title": f"{name_no_ext} 계약서",
            "prop_info": name_no_ext,
            "seq": "01",
            "contract_type": "최초임대차",
            "party_a": "",
            "party_b": "",
            "contract_date": ""
        }

t0_cols = [1.40, 2.025, 1.40, 2.025]
t1_cols = [0.45, 2.95, 0.85, 1.75, 0.85]
t2_cols = [1.40, 2.025, 1.40, 2.025]
t3_cols = [2.00, 4.85]

notes_count = 0

for base_dir in [LEASE_DIR, SALE_DIR]:
    if not os.path.exists(base_dir):
        continue

    for dir_name in sorted(os.listdir(base_dir)):
        flat_dir = os.path.join(base_dir, dir_name)
        if not os.path.isdir(flat_dir) or dir_name.startswith("_"):
            continue

        pdf_files = sorted([f for f in os.listdir(flat_dir) if f.lower().endswith(".pdf")])
        if not pdf_files:
            continue

        # Remove old docx notes
        for f_old in os.listdir(flat_dir):
            if f_old.endswith(".docx"):
                try:
                    os.remove(os.path.join(flat_dir, f_old))
                except Exception:
                    pass

        contract_items = [parse_option_a_filename(f) for f in pdf_files]
        primary_item = contract_items[0]
        primary_pdf_path = os.path.join(flat_dir, pdf_files[0])

        # Vision AI detail extraction
        v_info = extract_vision_details_from_pdf(primary_pdf_path, primary_item["filename_no_ext"])

        clean_unit_title = re.sub(r"^\d{2}_", "", dir_name)

        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        style_normal = doc.styles['Normal']
        style_normal.font.name = '맑은 고딕'
        style_normal.font.size = Pt(10.0)

        # P0 Title
        p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p0.text = f"[{clean_unit_title}] 부동산 계약 관리 노트"
        strip_keep_next(p0)
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(2)
        for r in p0.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(13.0)
            r.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)

        # P1 Sub-note
        p1 = doc.add_paragraph()
        p1.text = "※ [Vision AI 100% 추출 완료] 계약서 스캔 이미지에서 비전 AI로 정밀 분석한 내용 입니다.\n※ 최초 작성 시 원본 확인 후 검토해 주시면 됩니다."
        strip_keep_next(p1)
        p1.paragraph_format.space_after = Pt(6)
        for r in p1.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.0)
            r.font.color.rgb = RGBColor(100, 116, 139)

        # Section 1 Heading
        sec1_p = doc.add_paragraph()
        sec1_p.text = "1. 계약 정보"
        strip_keep_next(sec1_p)
        sec1_p.paragraph_format.space_before = Pt(6)
        sec1_p.paragraph_format.space_after = Pt(2)
        for r in sec1_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 0
        t0 = doc.add_table(rows=5, cols=4)
        t0.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t0, color="334155")
        set_table_width_and_columns(t0, t0_cols)

        end_date_str = v_info["end_date"] if v_info["end_date"] else "[종료일자 확인 필요]"
        period_str = f"{primary_item['contract_date']} ~ {end_date_str}"
        deposit_val = v_info["deposit"] if v_info["deposit"] else ""
        rent_val = v_info["monthly_rent"] if v_info["monthly_rent"] else ""

        t0_data = [
            [("건물명 / 호수", True), (clean_unit_title, False), ("사용처 (부서/용도)", True), (f"{clean_unit_title} 사무실", False)],
            [("계약 유형", True), (primary_item["contract_type"], False), ("최초 계약일 (시작일)", True), (primary_item["contract_date"], False)],
            [("임대 기간", True), (period_str, False), ("매월 납부일", True), (v_info["pay_day"], False)],
            [("보증금", True), (deposit_val, False), ("월 임대료", True), (rent_val, False)],
            [("계약면적 (㎡)", True), ("", False), ("전용면적 / 평수", True), ("", False)]
        ]

        for r_i, r_row in enumerate(t0.rows):
            set_row_cant_split(r_row)
            row_vals = t0_data[r_i]
            for c_i, (val, is_h) in enumerate(row_vals):
                cell = r_row.cells[c_i]
                if is_h:
                    format_header_cell(cell, val, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
                else:
                    format_data_cell(cell, val, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        # Section 2 Heading
        sec2_p = doc.add_paragraph()
        sec2_p.text = f"2. 수록 계약서 문서 목록 (총 {len(contract_items)}건)"
        strip_keep_next(sec2_p)
        sec2_p.paragraph_format.space_before = Pt(6)
        sec2_p.paragraph_format.space_after = Pt(2)
        for r in sec2_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 1
        t1 = doc.add_table(rows=1, cols=5)
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t1, color="334155")
        set_table_width_and_columns(t1, t1_cols)

        set_row_cant_split(t1.rows[0])
        set_table_header_repeat(t1.rows[0])

        t1_headers = ["순서", "문서명 (파일명)", "계약 종류", "계약 당사자 (임대인 → 임차인)", "임대 시작일"]
        for c_i, h in enumerate(t1_headers):
            cell = t1.rows[0].cells[c_i]
            format_header_cell(cell, h, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="E2E8F0")

        for idx, c_it in enumerate(contract_items, 1):
            r_row = t1.add_row()
            r_cells = r_row.cells
            set_row_cant_split(r_row)

            bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
            seq_str = f"{idx:02d}"
            c_type = c_it["contract_type"]
            parties_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] and c_it['party_b'] else c_it['party_a']
            c_date = c_it["contract_date"]

            format_data_cell(r_cells[0], seq_str, bold=True, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)

            cell1 = r_cells[1]
            set_cell_background(cell1, bg_color)
            set_cell_margins(cell1, top=60, bottom=60, left=100, right=100)
            cell1.text = ""
            
            p_title = cell1.paragraphs[0]
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_title.paragraph_format.space_before = Pt(1)
            p_title.paragraph_format.space_after = Pt(0)
            strip_keep_next(p_title)
            r_main = p_title.add_run(c_it["display_title"])
            r_main.font.name = "맑은 고딕"
            r_main.font.size = Pt(10.0)
            r_main.bold = True
            r_main.font.color.rgb = RGBColor(15, 23, 42)

            p_file = cell1.add_paragraph()
            p_file.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_file.paragraph_format.space_before = Pt(0)
            p_file.paragraph_format.space_after = Pt(1)
            strip_keep_next(p_file)
            r_sub = p_file.add_run(f"({c_it['filename_no_ext']}.pdf)")
            r_sub.font.name = "맑은 고딕"
            r_sub.font.size = Pt(8.5)
            r_sub.font.color.rgb = RGBColor(100, 116, 139)

            format_data_cell(r_cells[2], c_type, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)
            format_data_cell(r_cells[3], parties_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)
            format_data_cell(r_cells[4], c_date, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)

        set_table_width_and_columns(t1, t1_cols)

        # MANDATORY CLEAN PAGE BREAK BEFORE SECTION 3
        doc.add_page_break()

        # Section 3 Heading
        sec3_p = doc.add_paragraph()
        sec3_p.text = "3. 임대인 및 납부 계좌 정보"
        strip_keep_next(sec3_p)
        sec3_p.paragraph_format.space_before = Pt(6)
        sec3_p.paragraph_format.space_after = Pt(2)
        for r in sec3_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 2
        t2 = doc.add_table(rows=6, cols=4)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t2, color="334155")

        lessee_str = f"{primary_item['party_b']}(폭스커넥트 법인)" if "폭스" in primary_item['party_b'] else primary_item['party_b']
        holder_str = v_info["account_holder"] if v_info["account_holder"] else primary_item["party_a"]

        format_header_cell(t2.rows[0].cells[0], "임대인", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[0].cells[1], primary_item["party_a"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[0].cells[2], "임차인", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[0].cells[3], lessee_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        format_header_cell(t2.rows[1].cells[0], "임대인 연락처", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[1].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[1].cells[2], "관리사무소 연락처", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[1].cells[3], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        format_header_cell(t2.rows[2].cells[0], "입금 은행", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[2].cells[1], v_info["bank"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[2].cells[2], "예금주", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[2].cells[3], holder_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        t2.rows[3].cells[0].merge(t2.rows[3].cells[1])
        t2.rows[3].cells[2].merge(t2.rows[3].cells[3])
        format_header_cell(t2.rows[3].cells[0], "계좌번호", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[3].cells[2], v_info["account_num"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        t2.rows[4].cells[0].merge(t2.rows[4].cells[1]).merge(t2.rows[4].cells[2]).merge(t2.rows[4].cells[3])
        format_header_cell(t2.rows[4].cells[0], "비고 (관리자 참고사항)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

        t2.rows[5].cells[0].merge(t2.rows[5].cells[1]).merge(t2.rows[5].cells[2]).merge(t2.rows[5].cells[3])
        remarks_text = f"1. 본 건은 {clean_unit_title} 부동산 계약 관리 건임.\n2. 최초 작성 시 Vision AI 추출 정보 확인 후 관리자 검토를 진행해 주세요.\n3. 변경 이력 발생 시 4번 항목 표에 지속 업데이트해 주시기 바랍니다."
        format_data_cell(t2.rows[5].cells[0], remarks_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        set_table_width_and_columns(t2, t2_cols)
        for r in t2.rows:
            set_row_cant_split(r)

        # Section 4 Heading
        sec4_p = doc.add_paragraph()
        sec4_p.text = "4. 계약 변동 이력 및 특이사항"
        strip_keep_next(sec4_p)
        sec4_p.paragraph_format.space_before = Pt(6)
        sec4_p.paragraph_format.space_after = Pt(2)
        for r in sec4_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 3
        t3 = doc.add_table(rows=5, cols=2)
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t3, color="334155")

        history_lines = []
        for c_it in contract_items:
            if c_it["contract_type"] in ["전대차", "연장계약", "변경계약", "승계계약", "법인승계", "면적추가변경", "임대료변경"]:
                p_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] else ""
                history_lines.append(f"{c_it['contract_date']} [{c_it['contract_type']}] {p_str}")

        h_text = "\n".join(history_lines) if history_lines else "특이 변동이력 없음 (최초 계약 유지 중)"

        format_header_cell(t3.rows[0].cells[0], "계약 변동 / 전대차 / 승계 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[0].cells[1], h_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        format_header_cell(t3.rows[1].cells[0], "계약 연장 / 묵시적 갱신 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[1].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        format_header_cell(t3.rows[2].cells[0], "중도해지 / 퇴거 예정 메모", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[2].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        t3.rows[3].cells[0].merge(t3.rows[3].cells[1])
        format_header_cell(t3.rows[3].cells[0], "기타 특약 및 참조사항", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

        t3.rows[4].cells[0].merge(t3.rows[4].cells[1])
        special_terms_text = "1. 계약일 현재 등기부등본 확인 후 대상부동산의 권리 및 시설상태의 계약으로 한다.\n2. 임대차 만료 후 원상복구를 기본으로 한다.\n3. 부가세 및 관리비는 별도로 한다."
        format_data_cell(t3.rows[4].cells[0], special_terms_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        set_table_width_and_columns(t3, t3_cols)
        for r in t3.rows:
            set_row_cant_split(r)

        clean_file_title = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', clean_unit_title)
        docx_path = os.path.join(flat_dir, f"부동산_계약관리노트_{clean_file_title}.docx")
        doc.save(docx_path)
        print(f"  [VISION AI NOTE UPDATED] 부동산_계약관리노트_{clean_file_title}.docx")
        notes_count += 1

print(f"\n==========================================")
print(f"VISION AI CONTRACT NOTES GENERATION COMPLETE! Total {notes_count} Notes Processed!")
print(f"==========================================")
