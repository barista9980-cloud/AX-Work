import os
import re
import json
import fitz
import numpy as np
from PIL import Image
from rapidocr_onnxruntime import RapidOCR
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

DOWNLOAD_DIR = r"C:\Users\User\Downloads\drive-download-20260723T054808Z-1-001"
TARGET_ROOT = r"G:\내 드라이브\[FoxConnect]\[총무]업무\01_부동산_자산관리\01_임대차_전대차계약"
JSON_PATH = r"C:\Users\User\OneDrive\바탕 화면\업무_AX\real_estate_parsed_catalog.json"

ocr_engine = RapidOCR()

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def ocr_extract_pdf(pdf_path):
    deposit, rent, period, bank, account, phone, area = "", "", "", "", "", "", ""
    try:
        doc = fitz.open(pdf_path)
        ocr_texts = []
        for p_idx in range(min(3, len(doc))):
            pix = doc[p_idx].get_pixmap(dpi=250)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            results, _ = ocr_engine(np.array(img))
            if results:
                for res in results:
                    ocr_texts.append(res[1])

        combined_text = " ".join(ocr_texts)

        # 1. Bank Name
        for b_name in ["국민은행", "신한은행", "하나은행", "우리은행", "기업은행", "농협", "새마을금고", "부산은행", "대구은행"]:
            if b_name in combined_text:
                bank = b_name
                break

        # 2. Account Number
        acc_match = re.search(r'(\d{3,6}[\-\s]\d{2,6}[\-\s]\d{3,7})', combined_text)
        if acc_match:
            account = acc_match.group(1).strip()
            # Infer bank if account starts with 415 -> 국민은행
            if not bank and account.startswith("415"):
                bank = "국민은행"

        # 3. Financials
        amounts = re.findall(r'(\d{1,3}(?:[\.\,]\d{3})+)', combined_text)
        nums = []
        for a in amounts:
            clean_num = int(re.sub(r'[^\d]', '', a))
            if clean_num >= 10000:
                nums.append(clean_num)
        
        nums = sorted(list(set(nums)), reverse=True)
        if len(nums) >= 2:
            deposit = f"{nums[0]:,} 원"
            rent = f"{nums[1]:,} 원"
        elif len(nums) == 1:
            deposit = f"{nums[0]:,} 원"

        # 4. Phone numbers
        phone_match = re.findall(r'(0\d{1,2}[\-\s]\d{3,4}[\-\s]\d{4})', combined_text)
        if phone_match:
            phone = ", ".join(list(set(phone_match)))

        # 5. Area
        area_match = re.search(r'(\d{2,4}\.\d{1,2}|\d{2,4})\s*(㎡|평|m2)', combined_text)
        if area_match:
            area = f"{area_match.group(1)} ㎡"

        # 6. Period
        dates = re.findall(r'(20\d{2}[\.\-\s년]\s*\d{1,2}[\.\-\s월]\s*\d{1,2})', combined_text)
        if len(dates) >= 2:
            period = f"{dates[0]} ~ {dates[1]}"

    except Exception as e:
        print(f"OCR Error for {pdf_path}: {e}")

    return {
        "deposit": deposit,
        "rent": rent,
        "period": period,
        "bank": bank,
        "account": account,
        "phone": phone,
        "area": area
    }

def get_building_and_unit(region, prop_info):
    if "대륭포스트타워6차" in prop_info:
        building = f"{region}_대륭포스트타워6차"
        unit = "402_403호" if ("402" in prop_info or "403" in prop_info and "1510" not in prop_info) else ("1510호" if "1510" in prop_info else prop_info)
    elif "도곡로1길23" in prop_info:
        building = f"{region}_도곡로1길23"
        unit = prop_info.replace("도곡로1길23_", "")
    elif "GIDC" in prop_info:
        building = f"{region}_GIDC"
        if "1212" in prop_info:
            unit = "1212호"
        elif "1213" in prop_info:
            unit = "1213호"
        else:
            unit = "1214_1215호"
    elif "골프존" in prop_info:
        building = f"{region}_골프존"
        if "104" in prop_info:
            unit = "104호"
        elif "206" in prop_info:
            unit = "206호"
        else:
            unit = "204호"
    elif "스마트시티상가" in prop_info:
        building = f"{region}_스마트시티상가"
        unit = prop_info.replace("스마트시티상가_", "")
    elif "스마트시티" in prop_info:
        building = f"{region}_스마트시티"
        unit = prop_info.replace("스마트시티_", "")
    elif "하우스디어반" in prop_info:
        building = f"{region}_하우스디어반"
        unit = prop_info.replace("하우스디어반_", "")
    elif "_" in prop_info:
        parts = prop_info.split("_")
        building = f"{region}_{parts[0]}"
        unit = "_".join(parts[1:]) if len(parts) > 1 else "본건"
    else:
        building = f"{region}_{prop_info}"
        unit = "본건"

    return building, unit

def create_clean_docx(filepath, building_name, unit_name, items):
    doc = docx.Document()
    
    title = doc.add_heading(f"🏢 [{building_name}] {unit_name} 부동산 계약 관리 노트", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("※ [AI 파싱 완료] 🤖 표시는 계약서 PDF에서 자동 판독된 값이므로 내용이 맞는지 확인만 해주세요.\n※ [관리자 작성 필요]로 표시된 미입력 항목만 확인 후 작성해 주시면 됩니다.")
    p.runs[0].font.size = Pt(9.5)
    p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    primary_item = items[0]
    pdf_path = os.path.join(DOWNLOAD_DIR, primary_item["filename"])
    ocr_res = ocr_extract_pdf(pdf_path) if os.path.exists(pdf_path) else {}

    lessor = primary_item.get("party_a_lessor", "")
    lessee = primary_item.get("party_b_lessee", "")
    c_date = primary_item.get("contract_date", "")
    c_type = primary_item.get("contract_type", "")
    lessee_display = f"{lessee} (폭스커넥트 법인)" if "폭스" in lessee else lessee

    # Value or Manager tag logic
    dep_val = f"🤖 {ocr_res.get('deposit')}" if ocr_res.get('deposit') else "[관리자 작성 필요] (예시: 30,000,000 원)"
    rent_val = f"🤖 {ocr_res.get('rent')}" if ocr_res.get('rent') else "[관리자 작성 필요] (예시: 2,500,000 원)"
    period_val = f"🤖 {ocr_res.get('period')}" if ocr_res.get('period') else "[관리자 작성 필요] (예시: 2024-03-01 ~ 2026-02-28)"
    bank_val = f"🤖 {ocr_res.get('bank')}" if ocr_res.get('bank') else "[관리자 작성 필요] (예시: 국민은행)"
    account_val = f"🤖 {ocr_res.get('account')}" if ocr_res.get('account') else "[관리자 작성 필요] (예시: 123-456-7890)"
    phone_val = f"🤖 {ocr_res.get('phone')}" if ocr_res.get('phone') else "[관리자 작성 필요] (예시: 010-1234-5678)"
    area_val = f"🤖 {ocr_res.get('area')}" if ocr_res.get('area') else "[관리자 작성 필요] (예시: 150.5 ㎡)"

    doc.add_heading("1. 호수 및 주 계약 정보", level=2)
    t1 = doc.add_table(rows=5, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    data_t1 = [
        [("건물명 / 호수", f"{building_name} {unit_name}"), ("사용처 (부서/용도)", "[관리자 작성 필요] (예시: AX사업팀 본사 / 삭제 후 기입)")],
        [("주 계약 유형", f"🤖 {c_type}" if c_type else "임대차"), ("최초 계약일", f"🤖 {c_date}" if c_date else "[관리자 작성 필요]")],
        [("임대 기간", period_val), ("매월 납부일", "[관리자 작성 필요] (예시: 매월 25일 후불)")],
        [("보증금", dep_val), ("월 임대료", rent_val)],
        [("계약면적 (㎡)", area_val), ("전용면적 / 평수", "[관리자 작성 필요] (예시: 99.2 ㎡ / 30평)")]
    ]
    fill_table(t1, data_t1)

    # 2. All Related Contracts List
    doc.add_heading("2. 수록 계약서 문서 목록 (총 %d건)" % len(items), level=2)
    t_docs = doc.add_table(rows=len(items) + 1, cols=4)
    t_docs.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = t_docs.rows[0].cells
    hdr_cells[0].text = "구분/순서"
    hdr_cells[1].text = "계약 종류"
    hdr_cells[2].text = "계약 당사자 (임대인 - 임차인)"
    hdr_cells[3].text = "계약일"
    for c in hdr_cells:
        set_cell_background(c, "E6ECF5")
        
    for r_i, it in enumerate(items, 1):
        r_cells = t_docs.rows[r_i].cells
        r_cells[0].text = f"문서 {r_i:02d}"
        r_cells[1].text = it.get("contract_type", "")
        p_a = it.get("party_a_lessor", "")
        p_b = it.get("party_b_lessee", "")
        r_cells[2].text = f"{p_a} → {p_b}"
        r_cells[3].text = it.get("contract_date", "")

    # 3. Bank & Contacts
    doc.add_heading("3. 당사자 및 납부 계좌 정보", level=2)
    t2 = doc.add_table(rows=4, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    data_t2 = [
        [("주 임차인", f"🤖 {lessee_display}"), ("주 임대인", f"🤖 {lessor}")],
        [("임대인 연락처", phone_val), ("관리사무소 연락처", "[관리자 작성 필요] (예시: 02-1234-5678)")],
        [("입금 은행", bank_val), ("예금주", f"🤖 {lessor}")],
        [("계좌번호", account_val), ("비고", "특이사항 없음")]
    ]
    fill_table(t2, data_t2)

    # 4. Lifecycle History
    doc.add_heading("4. 계약 변동 이력 및 특이사항 (연 1~2회 업데이트)", level=2)
    t3 = doc.add_table(rows=4, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sublease_notes = []
    for it in items:
        ctype = it.get("contract_type", "")
        cdate = it.get("contract_date", "")
        pa = it.get("party_a_lessor", "")
        pb = it.get("party_b_lessee", "")
        if ctype in ["전대차", "연장계약", "변경계약", "법인승계"]:
            sublease_notes.append(f"{cdate} [{ctype}] {pa} → {pb}")
            
    history_text = f"🤖 " + "\n🤖 ".join(sublease_notes) if sublease_notes else "특이 변동이력 없음 (최초 계약 유지 중)"
    
    data_t3 = [
        [("계약 변동 / 전대차 / 승계 이력", history_text)],
        [("계약 연장 / 묵시적 갱신 이력", "[관리자 작성 필요] (작성 예시: 2026-03-01 묵시적 갱신 2년 연장)")],
        [("중도해지 / 퇴거 예정 메모", "[관리자 작성 필요] (작성 예시: 중도해지 사유 및 보증금 반환일)")],
        [("기타 특약 및 참조사항", "[관리자 작성 필요] (작성 예시: 제소전화해 완료 / 주차 지원)")]
    ]
    fill_table_2col(t3, data_t3)

    doc.save(filepath)

def fill_table(table, data):
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        row.cells[0].text = row_data[0][0]
        row.cells[1].text = row_data[0][1]
        row.cells[2].text = row_data[1][0]
        row.cells[3].text = row_data[1][1]
        set_cell_background(row.cells[0], "F2F4F7")
        set_cell_background(row.cells[2], "F2F4F7")

def fill_table_2col(table, data):
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        row.cells[0].text = row_data[0][0]
        row.cells[1].text = row_data[0][1]
        set_cell_background(row.cells[0], "F2F4F7")

def main():
    if not os.path.exists(JSON_PATH):
        print("Catalog JSON missing.")
        return

    with open(JSON_PATH, "r", encoding="utf-8") as f:
        catalog = json.load(f)

    buildings = {}
    for item in catalog:
        region = item.get("region", "기타")
        prop = item.get("property_info", "기타물건")
        b_name, u_name = get_building_and_unit(region, prop)
        
        if b_name not in buildings:
            buildings[b_name] = {}
        if u_name not in buildings[b_name]:
            buildings[b_name][u_name] = []
            
        buildings[b_name][u_name].append(item)

    print(f"Updating clean docx notes for {len(buildings)} buildings...")

    for b_idx, (b_name, units) in enumerate(buildings.items(), 1):
        b_folder_name = f"{b_idx:02d}_{b_name}"
        b_folder_name = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', b_folder_name)
        b_dir = os.path.join(TARGET_ROOT, b_folder_name)
        
        for u_idx, (u_name, items) in enumerate(units.items(), 1):
            u_folder_name = f"{u_name}"
            u_folder_name = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', u_folder_name)
            u_dir = os.path.join(b_dir, u_folder_name)
            
            if os.path.exists(u_dir):
                docx_path = os.path.join(u_dir, f"부동산_계약관리노트_{u_folder_name}.docx")
                create_clean_docx(docx_path, b_name, u_name, items)

        print(f"[{b_idx}/{len(buildings)}] Clean docx generated for: {b_folder_name}")

    print("\nAll docx templates successfully updated with clean OCR values!")

if __name__ == "__main__":
    main()
