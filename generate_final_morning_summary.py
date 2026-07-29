import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

notes_summary = []

for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    for f in files:
        if f.endswith(".docx") and "계약관리노트" in f:
            full_path = os.path.join(root, f)
            folder_name = os.path.basename(root)
            parent_name = os.path.basename(os.path.dirname(root))
            
            try:
                doc = docx.Document(full_path)
                t0 = doc.tables[0]
                contract_type = t0.rows[1].cells[1].text.strip()
                initial_date = t0.rows[1].cells[3].text.strip()
                period_str = t0.rows[2].cells[1].text.strip()
                pay_day = t0.rows[2].cells[3].text.strip()
                deposit_str = t0.rows[3].cells[1].text.strip()
                rent_str = t0.rows[3].cells[3].text.strip()
                
                t2 = doc.tables[2]
                lessor = t2.rows[0].cells[1].text.strip()
                lessee = t2.rows[0].cells[3].text.strip()
                
                notes_summary.append({
                    "building": parent_name,
                    "unit": folder_name,
                    "filename": f,
                    "lessor": lessor,
                    "lessee": lessee,
                    "contract_type": contract_type,
                    "period": period_str,
                    "deposit": deposit_str,
                    "rent": rent_str,
                    "pay_day": pay_day
                })
            except Exception as e:
                print(f"Error reading {f}: {e}")

print("=== FINAL OVERNIGHT PARSING RESULTS SUMMARY ===")
print(f"Total Folders Processed: {len(notes_summary)}")
for idx, item in enumerate(notes_summary, 1):
    print(f"{idx:02d}. [{item['building']}] {item['unit']}")
    print(f"    - 당사자: {item['lessor']} → {item['lessee']}")
    print(f"    - 임대기간: {item['period']}")
    print(f"    - 보증금: {item['deposit']} | 월세: {item['rent']} ({item['pay_day']})")
