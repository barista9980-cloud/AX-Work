import os
import re
import json
import docx
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
TEMPLATE_PATH = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_cell_text_and_style(cell, text, bold=False, font_size=9.0):
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(font_size)
            r.bold = bold

def parse_contract_filename(filename):
    name_no_ext = os.path.splitext(filename)[0]
    # Standard format: [지역_건물명]_순번_계약종류_[임대인-임차인]_(연월일)
    # Flexible regex:
    pattern = r"^(.*?)_(\d{2})_([^_]+)_\[(.*?)\]_\((\d{6})\)$"
    m = re.match(pattern, name_no_ext)
    if m:
        prop_info, seq, c_type, parties, raw_date = m.groups()
        party_list = parties.split('-') if '-' in parties else [parties]
        p_a = party_list[0] if len(party_list) > 0 else ""
        p_b = party_list[1] if len(party_list) > 1 else ""
        y = "20" + raw_date[:2]
        mon = raw_date[2:4]
        d = raw_date[4:6]
        f_date = f"{y}-{mon}-{d}"
        return {
            "filename_no_ext": name_no_ext,
            "prop_info": prop_info,
            "seq": seq,
            "contract_type": c_type,
            "party_a": p_a,
            "party_b": p_b,
            "contract_date": f_date
        }
    else:
        return {
            "filename_no_ext": name_no_ext,
            "prop_info": name_no_ext,
            "seq": "01",
            "contract_type": "최초임대차",
            "party_a": "",
            "party_b": "",
            "contract_date": ""
        }

processed_folders = 0
generated_notes = 0

print(f"Scanning {FOXCONNECT_ROOT} for property folders...")

for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    pdf_files = [f for f in files if f.lower().endswith(".pdf")]
    if not pdf_files:
        continue
    
    # Found a property folder containing PDFs
    print(f"\nProcessing Folder: {root}")
    
    # 1. Delete existing contract note files in this folder if present
    existing_docx = [f for f in files if f.lower().endswith(".docx") and ("계약관리노트" in f or "계약_노트" in f or "노트" in f)]
    for old_f in existing_docx:
        old_path = os.path.join(root, old_f)
        try:
            os.remove(old_path)
            print(f"  [DELETED OLD NOTE] {old_f}")
        except Exception as e:
            print(f"  [FAILED TO DELETE] {old_f}: {e}")

    # 2. Extract metadata from all PDFs in folder
    pdf_files.sort()
    contract_items = [parse_contract_filename(f) for f in pdf_files]
    
    primary_item = contract_items[0]
    folder_name = os.path.basename(root)
    parent_name = os.path.basename(os.path.dirname(root))
    
    b_name = parent_name
    u_name = folder_name
    
    # Create docx using 402_403호 template
    doc = docx.Document(TEMPLATE_PATH)
    
    # P0 Title
    doc.paragraphs[0].text = f"[{b_name}] {u_name} 부동산 계약 관리 노트"
    for r in doc.paragraphs[0].runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(15.0)
        r.bold = True
        
    # Table 0: Primary Contract Info
    t0 = doc.tables[0]
    t0_data = [
        [("건물명 / 호수", True), (f"{b_name} {u_name}", False), ("사용처 (부서/용도)", True), (f"{u_name} 사무실", False)],
        [("주 계약 유형", True), (primary_item["contract_type"], False), ("최초 계약일", True), (primary_item["contract_date"], False)],
        [("임대 기간", True), (f"{primary_item['contract_date']} ~ ", False), ("매월 납부일", True), ("", False)],
        [("보증금", True), ("", False), ("월 임대료", True), ("", False)],
        [("계약면적 (㎡)", True), ("", False), ("전용면적 / 평수", True), ("", False)]
    ]
    for r_i, r_data in enumerate(t0_data):
        for c_i, (val, is_h) in enumerate(r_data):
            cell = t0.rows[r_i].cells[c_i]
            set_cell_text_and_style(cell, val, bold=is_h)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # P3 Heading
    doc.paragraphs[3].text = f"2. 수록 계약서 문서 목록 (총 {len(contract_items)}건)"
    for r in doc.paragraphs[3].runs:
        r.font.name = "맑은 고딕"
        r.bold = True

    # Table 1: Docs List (5 cols: 순서 | 문서명 | 계약 종류 | 계약 당사자 | 계약일)
    t1 = doc.tables[1]
    while len(t1.rows) > 1:
        t1._tbl.remove(t1.rows[-1]._tr)
        
    t1_headers = ["순서", "문서명", "계약 종류", "계약 당사자 (임대인 - 임차인)", "계약일"]
    for c_i, h in enumerate(t1_headers):
        cell = t1.rows[0].cells[c_i]
        set_cell_text_and_style(cell, h, bold=True)
        set_cell_background(cell, "E6ECF5")
        
    for idx, c_it in enumerate(contract_items, 1):
        row_cells = t1.add_row().cells
        seq_str = f"{idx:02d}"
        f_name = c_it["filename_no_ext"]
        c_type = c_it["contract_type"]
        parties_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] and c_it['party_b'] else c_it['party_a']
        c_date = c_it["contract_date"]
        
        row_vals = [seq_str, f_name, c_type, parties_str, c_date]
        for c_i, v in enumerate(row_vals):
            set_cell_text_and_style(row_cells[c_i], v, bold=False)

    # Table 2: Parties & Bank Info
    t2 = doc.tables[2]
    lessee_str = f"{primary_item['party_b']} (폭스커넥트 관련 법인)" if "폭스" in primary_item['party_b'] else primary_item['party_b']
    t2_data = [
        [("주 임대인", True), (primary_item["party_a"], False), ("주 임차인", True), (lessee_str, False)],
        [("임대인 연락처", True), ("", False), ("관리사무소 연락처", True), ("", False)],
        [("입금 은행", True), ("", False), ("예금주", True), (primary_item["party_a"], False)],
        [("계좌번호", True), ("", False), ("비고", True), ("특이사항 없음", False)]
    ]
    for r_i, r_data in enumerate(t2_data):
        for c_i, (val, is_h) in enumerate(r_data):
            cell = t2.rows[r_i].cells[c_i]
            set_cell_text_and_style(cell, val, bold=is_h)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # Table 3: Lifecycle History
    t3 = doc.tables[3]
    history_lines = []
    for c_it in contract_items:
        if c_it["contract_type"] in ["전대차", "연장계약", "변경계약", "승계계약", "법인승계"]:
            p_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] else ""
            history_lines.append(f"{c_it['contract_date']}[{c_it['contract_type']}] {p_str}")
            
    h_text = "\n".join(history_lines) if history_lines else "특이 변동이력 없음 (최초 계약 유지 중)"
    
    t3_data = [
        [("계약 변동 / 전대차 / 승계 이력", True), (h_text, False)],
        [("계약 연장 / 묵시적 갱신 이력", True), ("", False)],
        [("중도해지 / 퇴거 예정 메모", True), ("", False)],
        [("기타 특약 및 참조사항", True), ("", False)]
    ]
    for r_i, r_data in enumerate(t3_data):
        for c_i, (val, is_h) in enumerate(r_data):
            cell = t3.rows[r_i].cells[c_i]
            set_cell_text_and_style(cell, val, bold=is_h)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # Save new contract note docx
    clean_u_name = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', u_name)
    new_docx_filename = f"부동산_계약관리노트_{clean_u_name}.docx"
    new_docx_path = os.path.join(root, new_docx_filename)
    
    try:
        doc.save(new_docx_path)
        print(f"  [CREATED NEW NOTE] {new_docx_filename}")
        generated_notes += 1
    except Exception as e:
        print(f"  [ERROR SAVING] {new_docx_filename}: {e}")

    processed_folders += 1

print(f"\n==========================================")
print(f"Finished! Total folders processed: {processed_folders}")
print(f"Total new contract notes generated: {generated_notes}")
print(f"==========================================")
