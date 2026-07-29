import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

docx_path = r"G:\내 드라이브\[FoxConnect]\[총무]업무\01_부동산_자산관리\01_임대차계약\01_판교_판교동612\부동산_계약관리노트_판교_판교동612.docx"

doc = docx.Document(docx_path)

# Update Table 0 Row 2 (임대 기간)
t0 = doc.tables[0]
t0.rows[2].cells[1].text = "2021-07-31 ~ 2023-07-30 (24개월)"

# Re-format cell
cell = t0.rows[2].cells[1]
p = cell.paragraphs[0]
p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs:
    r.font.name = "맑은 고딕"
    r.font.size = docx.shared.Pt(10.0)

doc.save(docx_path)
print("Successfully updated Pangyo lease period to: 2021-07-31 ~ 2023-07-30 (24개월)")
