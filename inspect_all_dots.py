import os
import docx

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

doc_file = None
for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    for f in files:
        if f.endswith(".docx") and "1510" in f:
            doc_file = os.path.join(root, f)
            break
    if doc_file:
        break

if doc_file:
    print(f"Inspecting file: {doc_file}")
    doc = docx.Document(doc_file)
    print("\n--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        print(f"P{i}: '{p.text}'")

    print("\n--- TABLES TEXT ---")
    for t_idx, t in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        for r_idx, r in enumerate(t.rows):
            for c_idx, c in enumerate(r.cells):
                print(f"  T{t_idx} R{r_idx} C{c_idx}: '{c.text}'")
else:
    print("File not found.")
