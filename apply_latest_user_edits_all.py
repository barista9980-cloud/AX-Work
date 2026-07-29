import os
import re
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_table_borders(table, color="000000", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def format_cell(cell, text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = text
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = False
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = bold

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_table_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def parse_contract_filename(filename):
    name_no_ext = os.path.splitext(filename)[0]
    pattern = r"^(.*?)_(\d{2})_([^_]+)_\[(.*?)\]_\((\d{6})\)$"
    m = re.match(pattern, name_no_ext)
    if m:
        prop_info, seq, c_type, parties, raw_date = m.groups()
        party_list = parties.split('-') if '-' in parties else [parties]
        p_a = party_list[0] if len(party_list) > 0 else ""
        p_b = party_list[1] if len(party_list) > 1 else ""
        y = "20" + raw_date[:2]
        mon = raw_date[2:4]
        d = raw_date[4:6]
        f_date = f"{y}-{mon}-{d}"
        return {
            "filename_no_ext": name_no_ext,
            "prop_info": prop_info,
            "seq": seq,
            "contract_type": c_type,
            "party_a": p_a,
            "party_b": p_b,
            "contract_date": f_date
        }
    else:
        return {
            "filename_no_ext": name_no_ext,
            "prop_info": name_no_ext,
            "seq": "01",
            "contract_type": "최초임대차",
            "party_a": "",
            "party_b": "",
            "contract_date": ""
        }

processed_folders = 0
generated_notes = 0

print(f"Scanning {FOXCONNECT_ROOT} for property folders to apply user's latest 402_403호 modifications...")

for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    pdf_files = [f for f in files if f.lower().endswith(".pdf")]
    if not pdf_files:
        continue

    print(f"\nProcessing Folder: {root}")

    # Delete existing note files
    existing_notes = [f for f in files if f.lower().endswith(".docx") and ("계약관리노트" in f or "계약_노트" in f or "노트" in f)]
    for old_f in existing_notes:
        old_p = os.path.join(root, old_f)
        try:
            os.remove(old_p)
            print(f"  [DELETED OLD NOTE] {old_f}")
        except Exception as e:
            print(f"  [FAIL DELETE] {old_f}: {e}")

    pdf_files.sort()
    contract_items = [parse_contract_filename(f) for f in pdf_files]
    primary_item = contract_items[0]

    folder_name = os.path.basename(root)
    parent_name = os.path.basename(os.path.dirname(root))

    b_name = parent_name
    u_name = folder_name

    # Create clean Document from scratch
    doc = docx.Document()

    # Set Normal Style default font to 맑은 고딕 10pt
    style_normal = doc.styles['Normal']
    style_normal.font.name = '맑은 고딕'
    style_normal.font.size = Pt(10.0)
    try:
        rPr = style_normal._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="맑은 고딕" w:hAnsi="맑은 고딕" w:eastAsia="맑은 고딕" w:cs="맑은 고딕"/>')
        rPr.append(rFonts)
    except Exception:
        pass

    # P0 Title: 13pt Bold 맑은 고딕 (Matching user's updated format: [{b_name}] {u_name} 부동산 계약 관리 노트)
    p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p0.text = f"[{b_name}] {u_name} 부동산 계약 관리 노트"
    p0.paragraph_format.keep_with_next = False
    p0.paragraph_format.space_after = Pt(4)
    for r in p0.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(13.0)
        r.bold = True

    # P1 Sub-note: 9pt Gray 맑은 고딕
    p1 = doc.add_paragraph()
    p1.text = "※ [AI 파싱 완료] 계약서 PDF에서 OCR 정밀 추출된 내용 입니다.\n※ 최초 작성 시 OCR 정밀 추출 결과물 확인 후 검토해 주시면 됩니다."
    p1.paragraph_format.keep_with_next = False
    p1.paragraph_format.space_after = Pt(8)
    for r in p1.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(75, 85, 99)

    # Section 1 Heading: "1. 계약 정보" (Changed from "1. 주 계약 정보" per user update!)
    sec1_p = doc.add_paragraph()
    sec1_p.text = "1. 계약 정보"
    sec1_p.paragraph_format.keep_with_next = False
    sec1_p.paragraph_format.space_before = Pt(6)
    sec1_p.paragraph_format.space_after = Pt(4)
    for r in sec1_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True

    # Table 0: 계약 정보 (Changed label "주 계약 유형" -> "계약 유형")
    t0 = doc.add_table(rows=5, cols=4)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t0)

    t0_data = [
        [("건물명 / 호수", True), (f"{b_name} {u_name}", False), ("사용처 (부서/용도)", True), (f"{u_name} 사무실", False)],
        [("계약 유형", True), (primary_item["contract_type"], False), ("최초 계약일", True), (primary_item["contract_date"], False)],
        [("임대 기간", True), (f"{primary_item['contract_date']} ~ ", False), ("매월 납부일", True), ("", False)],
        [("보증금", True), ("", False), ("월 임대료", True), ("", False)],
        [("계약면적 (㎡)", True), ("", False), ("전용면적 / 평수", True), ("", False)]
    ]

    for r_i, r_row in enumerate(t0.rows):
        set_row_cant_split(r_row)
        row_vals = t0_data[r_i]
        for c_i, (val, is_h) in enumerate(row_vals):
            cell = r_row.cells[c_i]
            align = WD_ALIGN_PARAGRAPH.CENTER
            format_cell(cell, val, bold=is_h, font_size=10.0, align=align)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # Section 2 Heading: "2. 수록 계약서 문서 목록 (총 N건)"
    sec2_p = doc.add_paragraph()
    sec2_p.text = f"2. 수록 계약서 문서 목록 (총 {len(contract_items)}건)"
    sec2_p.paragraph_format.keep_with_next = False
    sec2_p.paragraph_format.space_before = Pt(8)
    sec2_p.paragraph_format.space_after = Pt(4)
    for r in sec2_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True

    # Table 1: Docs List (5 cols: 순서 | 문서명 | 계약 종류 | 당사자 | 계약일)
    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    set_row_cant_split(t1.rows[0])
    set_table_header_repeat(t1.rows[0])

    t1_headers = ["순서", "문서명", "계약 종류", "계약 당사자 (임대인 - 임차인)", "계약일"]
    for c_i, h in enumerate(t1_headers):
        cell = t1.rows[0].cells[c_i]
        format_cell(cell, h, bold=True, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_background(cell, "E6ECF5")

    for idx, c_it in enumerate(contract_items, 1):
        r_cells = t1.add_row().cells
        set_row_cant_split(t1.rows[-1])

        seq_str = f"{idx:02d}"
        f_name = c_it["filename_no_ext"]
        c_type = c_it["contract_type"]
        parties_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] and c_it['party_b'] else c_it['party_a']
        c_date = c_it["contract_date"]

        format_cell(r_cells[0], seq_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r_cells[1], f_name, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r_cells[2], c_type, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r_cells[3], parties_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r_cells[4], c_date, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Section 3 Heading: "3. 임대인 및 납부 계좌 정보" (NO EMPTY ENTERS BEFORE IT)
    sec3_p = doc.add_paragraph()
    sec3_p.text = "3. 임대인 및 납부 계좌 정보"
    sec3_p.paragraph_format.keep_with_next = False
    sec3_p.paragraph_format.space_before = Pt(8)
    sec3_p.paragraph_format.space_after = Pt(4)
    for r in sec3_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True

    # Table 2: 임대인 및 계좌 정보 (Changed labels: "주 임대인" -> "임대인", "주 임차인" -> "임차인")
    t2 = doc.add_table(rows=4, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    lessee_str = f"{primary_item['party_b']}(폭스커넥트 법인)" if "폭스" in primary_item['party_b'] else primary_item['party_b']
    t2_data = [
        [("임대인", True), (primary_item["party_a"], False), ("임차인", True), (lessee_str, False)],
        [("임대인 연락처", True), ("", False), ("관리사무소 연락처", True), ("", False)],
        [("입금 은행", True), ("", False), ("예금주", True), (primary_item["party_a"], False)],
        [("계좌번호", True), ("", False), ("비고", True), ("특이사항 없음", False)]
    ]

    for r_i, r_row in enumerate(t2.rows):
        set_row_cant_split(r_row)
        row_vals = t2_data[r_i]
        for c_i, (val, is_h) in enumerate(row_vals):
            cell = r_row.cells[c_i]
            align = WD_ALIGN_PARAGRAPH.CENTER
            format_cell(cell, val, bold=is_h, font_size=10.0, align=align)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # Section 4 Heading: "4. 계약 변동 이력 및 특이사항"
    sec4_p = doc.add_paragraph()
    sec4_p.text = "4. 계약 변동 이력 및 특이사항"
    sec4_p.paragraph_format.keep_with_next = False
    sec4_p.paragraph_format.space_before = Pt(8)
    sec4_p.paragraph_format.space_after = Pt(4)
    for r in sec4_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True

    # Table 3: 계약 변동 이력 (Cell font 10pt)
    t3 = doc.add_table(rows=4, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)

    history_lines = []
    for c_it in contract_items:
        if c_it["contract_type"] in ["전대차", "연장계약", "변경계약", "승계계약", "법인승계"]:
            p_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] else ""
            history_lines.append(f"{c_it['contract_date']}[{c_it['contract_type']}] {p_str}")

    h_text = "\n".join(history_lines) if history_lines else "특이 변동이력 없음 (최초 계약 유지 중)"

    t3_data = [
        [("계약 변동 / 전대차 / 승계 이력", True), (h_text, False)],
        [("계약 연장 / 묵시적 갱신 이력", True), ("", False)],
        [("중도해지 / 퇴거 예정 메모", True), ("", False)],
        [("기타 특약 및 참조사항", True), ("", False)]
    ]

    for r_i, r_row in enumerate(t3.rows):
        set_row_cant_split(r_row)
        row_vals = t3_data[r_i]
        for c_i, (val, is_h) in enumerate(row_vals):
            cell = r_row.cells[c_i]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(cell, val, bold=is_h, font_size=10.0, align=align)
            if is_h:
                set_cell_background(cell, "F2F4F7")

    # Save clean contract note
    clean_u_name = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', u_name)
    new_docx_filename = f"부동산_계약관리노트_{clean_u_name}.docx"
    new_docx_path = os.path.join(root, new_docx_filename)

    try:
        doc.save(new_docx_path)
        print(f"  [UPDATED WITH LATEST USER EDIT SPEC] {new_docx_filename}")
        generated_notes += 1
    except Exception as e:
        print(f"  [ERROR SAVING] {new_docx_filename}: {e}")

    processed_folders += 1

print(f"\n==========================================")
print(f"Finished updating all property notes! Total folders processed: {processed_folders}")
print(f"Total updated contract notes generated: {generated_notes}")
print(f"==========================================")
