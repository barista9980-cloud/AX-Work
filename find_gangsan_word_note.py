import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
GANGSAN_DIR = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리\01_임대차계약\22_강남_도곡로1길23_지하1층,1층,2층,3층")

docx_files = [f for f in os.listdir(GANGSAN_DIR) if f.endswith(".docx")]
print("Found Docx files:", docx_files)

for f in docx_files:
    dp = os.path.join(GANGSAN_DIR, f)
    print(f"\n==========================================")
    print(f"Reading: {f}")
    print(f"==========================================")
    doc = docx.Document(dp)
    for p in doc.paragraphs:
        if p.text.strip():
            print("P:", p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            row_txt = [c.text.strip().replace("\n", " ") for c in r.cells]
            print("TBL:", " | ".join(row_txt))
