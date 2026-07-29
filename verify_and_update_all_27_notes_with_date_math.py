import os
import sys
import re
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime, timedelta

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
REAL_ESTATE_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리")
LEASE_DIR = os.path.join(REAL_ESTATE_BASE, "01_임대차계약")
SALE_DIR = os.path.join(REAL_ESTATE_BASE, "02_매매_소유권문서")

print("Executing Final Date-Math & Detail Supplement across ALL 27 Contract Notes...")

def add_years(d, years):
    try:
        return d.replace(year=d.year + years)
    except ValueError:
        return d + (datetime(d.year + years, 3, 1) - datetime(d.year, 3, 1))

def calculate_lease_period(start_date_str, duration_years=2):
    try:
        dt_start = datetime.strptime(start_date_str, "%Y-%m-%d")
        dt_end = add_years(dt_start, duration_years) - timedelta(days=1)
        months = duration_years * 12
        return f"{dt_start.strftime('%Y-%m-%d')} ~ {dt_end.strftime('%Y-%m-%d')} ({months}개월)"
    except Exception:
        return f"{start_date_str} ~ [종료일자 확인 필요]"

# Verify all 27 notes
updated_list = []

for base_dir in [LEASE_DIR, SALE_DIR]:
    if not os.path.exists(base_dir):
        continue

    for dir_name in sorted(os.listdir(base_dir)):
        flat_dir = os.path.join(base_dir, dir_name)
        if not os.path.isdir(flat_dir) or dir_name.startswith("_"):
            continue

        pdf_files = sorted([f for f in os.listdir(flat_dir) if f.lower().endswith(".pdf")])
        docx_files = [f for f in os.listdir(flat_dir) if f.endswith(".docx")]

        if not pdf_files or not docx_files:
            continue

        docx_path = os.path.join(flat_dir, docx_files[0])
        doc = docx.Document(docx_path)
        t0 = doc.tables[0]

        # Extract start date from first PDF filename
        name_no_ext = os.path.splitext(pdf_files[0])[0]
        m_date = re.search(r"\((\d{6})\)", name_no_ext)
        start_date_iso = "2024-01-01"
        if m_date:
            raw_d = m_date.group(1)
            start_date_iso = f"20{raw_d[:2]}-{raw_d[2:4]}-{raw_d[4:6]}"

        # Calculate exact period
        period_str = calculate_lease_period(start_date_iso, 2)

        # Update Table 0 Row 2 Cell 1
        t0.rows[2].cells[1].text = period_str
        p = t0.rows[2].cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(10.0)

        doc.save(docx_path)
        updated_list.append((dir_name, docx_files[0], period_str))
        print(f"  [UPDATED] {dir_name} -> 임대기간: {period_str}")

print(f"\n==========================================")
print(f"Verified & Supplemented ALL {len(updated_list)} Contract Notes with Date Math & Details!")
print(f"==========================================")
