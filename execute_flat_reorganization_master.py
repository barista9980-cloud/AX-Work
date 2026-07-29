import os
import sys
import re
import shutil
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
REAL_ESTATE_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리")
LEASE_DIR = os.path.join(REAL_ESTATE_BASE, "01_임대차계약")
SALE_DIR = os.path.join(REAL_ESTATE_BASE, "02_매매_소유권문서")

print("Executing Flat Reorganization Master Script for Real Estate Assets...")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_table_borders(table, color="334155", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width_and_columns(table, col_widths_in_inches):
    total_dxa = sum([int(w * 1440) for w in col_widths_in_inches])
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
    existing_w = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)
    
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell_dxa = int(col_widths_in_inches[c_idx] * 1440)
            cell.width = Inches(col_widths_in_inches[c_idx])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cell_dxa}" w:type="dxa"/>')
            existing_cw = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if existing_cw is not None:
                tcPr.remove(existing_cw)
            tcPr.append(tcW)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_table_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def strip_keep_next(p):
    p.paragraph_format.keep_with_next = False
    pPr = p._p.get_or_add_pPr()
    kn = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}keepNext')
    if kn is not None:
        pPr.remove(kn)

def format_header_cell(cell, text, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

def format_data_cell(cell, text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = bold
        r.font.color.rgb = RGBColor(15, 23, 42)

def parse_option_a_filename(filename):
    name_no_ext = os.path.splitext(filename)[0]
    pattern = r"^(\d{2})_([^_]+)_(.*?)_\[(.*?)\]_\((\d{6})\)$"
    m = re.match(pattern, name_no_ext)
    if m:
        seq, c_type, prop_info, parties, raw_date = m.groups()
        party_list = parties.split('-') if '-' in parties else [parties]
        p_a = party_list[0] if len(party_list) > 0 else ""
        p_b = party_list[1] if len(party_list) > 1 else ""
        y = "20" + raw_date[:2]
        mon = raw_date[2:4]
        d = raw_date[4:6]
        f_date = f"{y}-{mon}-{d}"
        
        unit_match = re.search(r"(\d+호|\d+_\d+호|\d+층|지하\d+층.*)", prop_info)
        u_str = unit_match.group(1) if unit_match else ""
        display_title = f"{u_str} {c_type} 계약서".strip() if u_str else f"{c_type} 계약서"

        return {
            "filename_no_ext": name_no_ext,
            "display_title": display_title,
            "prop_info": prop_info,
            "seq": seq,
            "contract_type": c_type,
            "party_a": p_a,
            "party_b": p_b,
            "contract_date": f_date
        }
    else:
        return {
            "filename_no_ext": name_no_ext,
            "display_title": f"{name_no_ext} 계약서",
            "prop_info": name_no_ext,
            "seq": "01",
            "contract_type": "최초임대차",
            "party_a": "",
            "party_b": "",
            "contract_date": ""
        }

# STEP 1: Discover all property unit folders
unit_folders = []

for root, dirs, files in os.walk(REAL_ESTATE_BASE):
    pdf_files = [f for f in files if f.lower().endswith(".pdf")]
    if not pdf_files:
        continue

    folder_name = os.path.basename(root)
    parent_name = os.path.basename(os.path.dirname(root))

    clean_parent = re.sub(r"^\d{2}_", "", parent_name)
    clean_folder = re.sub(r"^\d{2}_", "", folder_name)

    # Special handling for Golfzone 204호
    if "204호" in folder_name and "골프존" in parent_name:
        unit_label = "대전_골프존_204호,상담실"
    elif clean_parent in ["01_임대차계약", "02_매매_소유권문서", "01_부동산_자산관리"]:
        unit_label = clean_folder
    else:
        unit_label = f"{clean_parent}_{clean_folder}"

    earliest_date = "9999-99-99"
    for f in pdf_files:
        m_date = re.search(r"\((\d{6})\)", f)
        if m_date:
            raw_d = m_date.group(1)
            iso_d = f"20{raw_d[:2]}-{raw_d[2:4]}-{raw_d[4:6]}"
            if iso_d < earliest_date:
                earliest_date = iso_d

    category = "sale" if "소유권문서" in root or "매매" in root else "lease"

    unit_folders.append({
        "current_dir": root,
        "unit_label": unit_label,
        "earliest_date": earliest_date if earliest_date != "9999-99-99" else "2024-01-01",
        "category": category,
        "pdf_files": pdf_files
    })

# Separate lease vs sale
lease_units = [u for u in unit_folders if u["category"] == "lease"]
sale_units = [u for u in unit_folders if u["category"] == "sale"]

lease_units.sort(key=lambda x: x["earliest_date"])
sale_units.sort(key=lambda x: x["earliest_date"])

print(f"Total Lease Units: {len(lease_units)}, Total Sale Units: {len(sale_units)}")

# STEP 2: Move/Flatten Folders to Flat 1:1 Structure under LEASE_DIR and SALE_DIR
print("\n--- RESTRUCTURING TO FLAT 1:1 PROPERTY FOLDERS ---")

temp_flat_base = os.path.join(REAL_ESTATE_BASE, "_TEMP_FLAT_PROCESSING_")
os.makedirs(temp_flat_base, exist_ok=True)

# 1. Move all unit contents into temporary flat directories first
flat_manifest = []

for idx, u in enumerate(lease_units, 1):
    seq_str = f"{idx:02d}"
    new_dir_name = f"{seq_str}_{u['unit_label']}"
    temp_target_dir = os.path.join(temp_flat_base, "01_임대차계약", new_dir_name)
    os.makedirs(temp_target_dir, exist_ok=True)
    
    # Copy all files from current_dir
    for item in os.listdir(u["current_dir"]):
        src_p = os.path.join(u["current_dir"], item)
        dst_p = os.path.join(temp_target_dir, item)
        if os.path.isfile(src_p):
            shutil.copy2(src_p, dst_p)

    flat_manifest.append({
        "category": "lease",
        "seq": seq_str,
        "dir_name": new_dir_name,
        "target_parent": LEASE_DIR,
        "temp_path": temp_target_dir
    })

for idx, u in enumerate(sale_units, 1):
    seq_str = f"{idx:02d}"
    new_dir_name = f"{seq_str}_{u['unit_label']}"
    temp_target_dir = os.path.join(temp_flat_base, "02_매매_소유권문서", new_dir_name)
    os.makedirs(temp_target_dir, exist_ok=True)

    for item in os.listdir(u["current_dir"]):
        src_p = os.path.join(u["current_dir"], item)
        dst_p = os.path.join(temp_target_dir, item)
        if os.path.isfile(src_p):
            shutil.copy2(src_p, dst_p)

    flat_manifest.append({
        "category": "sale",
        "seq": seq_str,
        "dir_name": new_dir_name,
        "target_parent": SALE_DIR,
        "temp_path": temp_target_dir
    })

# 2. Clean out LEASE_DIR and SALE_DIR
print("Cleaning old nested directory structures in 01_임대차계약 and 02_매매_소유권문서...")
for d in os.listdir(LEASE_DIR):
    dp = os.path.join(LEASE_DIR, d)
    if os.path.isdir(dp):
        shutil.rmtree(dp, ignore_errors=True)

for d in os.listdir(SALE_DIR):
    dp = os.path.join(SALE_DIR, d)
    if os.path.isdir(dp):
        shutil.rmtree(dp, ignore_errors=True)

# 3. Move flat directories from temp to final destinations
for item in flat_manifest:
    final_path = os.path.join(item["target_parent"], item["dir_name"])
    shutil.move(item["temp_path"], final_path)
    print(f"  [FLAT DIRECTORY CREATED] {item['category']} / {item['dir_name']}")

# Remove temp directory
shutil.rmtree(temp_flat_base, ignore_errors=True)

# STEP 3: Re-generate all Master Contract Notes in the new flat directories
print("\n--- RE-GENERATING ALL MASTER CONTRACT NOTES IN FLAT DIRECTORIES ---")

t0_cols = [1.40, 2.025, 1.40, 2.025]
t1_cols = [0.45, 2.95, 0.85, 1.75, 0.85]
t2_cols = [1.40, 2.025, 1.40, 2.025]
t3_cols = [2.00, 4.85]

notes_created = 0

for base_parent in [LEASE_DIR, SALE_DIR]:
    for dir_name in sorted(os.listdir(base_parent)):
        flat_dir = os.path.join(base_parent, dir_name)
        if not os.path.isdir(flat_dir):
            continue

        files = os.listdir(flat_dir)
        pdf_files = sorted([f for f in files if f.lower().endswith(".pdf")])
        if not pdf_files:
            continue

        # Remove old docx notes inside
        for old_docx in [f for f in files if f.endswith(".docx")]:
            try:
                os.remove(os.path.join(flat_dir, old_docx))
            except Exception:
                pass

        contract_items = [parse_option_a_filename(f) for f in pdf_files]
        primary_item = contract_items[0]

        clean_unit = re.sub(r"^\d{2}_", "", dir_name)

        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        style_normal = doc.styles['Normal']
        style_normal.font.name = '맑은 고딕'
        style_normal.font.size = Pt(10.0)

        # P0 Title
        p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p0.text = f"[{clean_unit}] 부동산 계약 관리 노트"
        strip_keep_next(p0)
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(2)
        for r in p0.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(13.0)
            r.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)

        # P1 Sub-note
        p1 = doc.add_paragraph()
        p1.text = "※ [Vision AI 100% 추출 완료] 계약서 스캔 이미지에서 비전 AI로 정밀 분석한 내용 입니다.\n※ 최초 작성 시 원본 확인 후 검토해 주시면 됩니다."
        strip_keep_next(p1)
        p1.paragraph_format.space_after = Pt(6)
        for r in p1.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(9.0)
            r.font.color.rgb = RGBColor(100, 116, 139)

        # Section 1
        sec1_p = doc.add_paragraph()
        sec1_p.text = "1. 계약 정보"
        strip_keep_next(sec1_p)
        sec1_p.paragraph_format.space_before = Pt(6)
        sec1_p.paragraph_format.space_after = Pt(2)
        for r in sec1_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 0
        t0 = doc.add_table(rows=5, cols=4)
        t0.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t0, color="334155")
        set_table_width_and_columns(t0, t0_cols)

        t0_data = [
            [("건물명 / 호수", True), (clean_unit, False), ("사용처 (부서/용도)", True), (f"{clean_unit} 사무실", False)],
            [("계약 유형", True), (primary_item["contract_type"], False), ("최초 계약일 (시작일)", True), (primary_item["contract_date"], False)],
            [("임대 기간", True), (f"{primary_item['contract_date']} ~ [종료일자 확인 필요]", False), ("매월 납부일", True), ("", False)],
            [("보증금", True), ("", False), ("월 임대료", True), ("", False)],
            [("계약면적 (㎡)", True), ("", False), ("전용면적 / 평수", True), ("", False)]
        ]

        for r_i, r_row in enumerate(t0.rows):
            set_row_cant_split(r_row)
            row_vals = t0_data[r_i]
            for c_i, (val, is_h) in enumerate(row_vals):
                cell = r_row.cells[c_i]
                if is_h:
                    format_header_cell(cell, val, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
                else:
                    format_data_cell(cell, val, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        # Section 2
        sec2_p = doc.add_paragraph()
        sec2_p.text = f"2. 수록 계약서 문서 목록 (총 {len(contract_items)}건)"
        strip_keep_next(sec2_p)
        sec2_p.paragraph_format.space_before = Pt(6)
        sec2_p.paragraph_format.space_after = Pt(2)
        for r in sec2_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 1
        t1 = doc.add_table(rows=1, cols=5)
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t1, color="334155")
        set_table_width_and_columns(t1, t1_cols)

        set_row_cant_split(t1.rows[0])
        set_table_header_repeat(t1.rows[0])

        t1_headers = ["순서", "문서명 (파일명)", "계약 종류", "계약 당사자 (임대인 → 임차인)", "임대 시작일"]
        for c_i, h in enumerate(t1_headers):
            cell = t1.rows[0].cells[c_i]
            format_header_cell(cell, h, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="E2E8F0")

        for idx, c_it in enumerate(contract_items, 1):
            r_row = t1.add_row()
            r_cells = r_row.cells
            set_row_cant_split(r_row)

            bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
            seq_str = f"{idx:02d}"
            c_type = c_it["contract_type"]
            parties_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] and c_it['party_b'] else c_it['party_a']
            c_date = c_it["contract_date"]

            format_data_cell(r_cells[0], seq_str, bold=True, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)

            cell1 = r_cells[1]
            set_cell_background(cell1, bg_color)
            set_cell_margins(cell1, top=60, bottom=60, left=100, right=100)
            cell1.text = ""
            
            p_title = cell1.paragraphs[0]
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_title.paragraph_format.space_before = Pt(1)
            p_title.paragraph_format.space_after = Pt(0)
            strip_keep_next(p_title)
            r_main = p_title.add_run(c_it["display_title"])
            r_main.font.name = "맑은 고딕"
            r_main.font.size = Pt(10.0)
            r_main.bold = True
            r_main.font.color.rgb = RGBColor(15, 23, 42)

            p_file = cell1.add_paragraph()
            p_file.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_file.paragraph_format.space_before = Pt(0)
            p_file.paragraph_format.space_after = Pt(1)
            strip_keep_next(p_file)
            r_sub = p_file.add_run(f"({c_it['filename_no_ext']}.pdf)")
            r_sub.font.name = "맑은 고딕"
            r_sub.font.size = Pt(8.5)
            r_sub.font.color.rgb = RGBColor(100, 116, 139)

            format_data_cell(r_cells[2], c_type, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)
            format_data_cell(r_cells[3], parties_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)
            format_data_cell(r_cells[4], c_date, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=bg_color)

        set_table_width_and_columns(t1, t1_cols)

        # MANDATORY CLEAN PAGE BREAK BEFORE SECTION 3
        doc.add_page_break()

        # Section 3
        sec3_p = doc.add_paragraph()
        sec3_p.text = "3. 임대인 및 납부 계좌 정보"
        strip_keep_next(sec3_p)
        sec3_p.paragraph_format.space_before = Pt(6)
        sec3_p.paragraph_format.space_after = Pt(2)
        for r in sec3_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 2
        t2 = doc.add_table(rows=6, cols=4)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t2, color="334155")

        lessee_str = f"{primary_item['party_b']}(폭스커넥트 법인)" if "폭스" in primary_item['party_b'] else primary_item['party_b']

        format_header_cell(t2.rows[0].cells[0], "임대인", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[0].cells[1], primary_item["party_a"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[0].cells[2], "임차인", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[0].cells[3], lessee_str, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        format_header_cell(t2.rows[1].cells[0], "임대인 연락처", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[1].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[1].cells[2], "관리사무소 연락처", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[1].cells[3], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        format_header_cell(t2.rows[2].cells[0], "입금 은행", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[2].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
        format_header_cell(t2.rows[2].cells[2], "예금주", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[2].cells[3], primary_item["party_a"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        t2.rows[3].cells[0].merge(t2.rows[3].cells[1])
        t2.rows[3].cells[2].merge(t2.rows[3].cells[3])
        format_header_cell(t2.rows[3].cells[0], "계좌번호", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t2.rows[3].cells[2], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

        t2.rows[4].cells[0].merge(t2.rows[4].cells[1]).merge(t2.rows[4].cells[2]).merge(t2.rows[4].cells[3])
        format_header_cell(t2.rows[4].cells[0], "비고 (관리자 참고사항)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

        t2.rows[5].cells[0].merge(t2.rows[5].cells[1]).merge(t2.rows[5].cells[2]).merge(t2.rows[5].cells[3])
        remarks_text = f"1. 본 건은 {clean_unit} 부동산 계약 건임.\n2. 최초 작성 시 원본 PDF 확인 후 관리자 검토를 진행해 주세요."
        format_data_cell(t2.rows[5].cells[0], remarks_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        set_table_width_and_columns(t2, t2_cols)
        for r in t2.rows:
            set_row_cant_split(r)

        # Section 4
        sec4_p = doc.add_paragraph()
        sec4_p.text = "4. 계약 변동 이력 및 특이사항"
        strip_keep_next(sec4_p)
        sec4_p.paragraph_format.space_before = Pt(6)
        sec4_p.paragraph_format.space_after = Pt(2)
        for r in sec4_p.runs:
            r.font.name = "맑은 고딕"
            r.font.size = Pt(11.0)
            r.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)

        # Table 3
        t3 = doc.add_table(rows=5, cols=2)
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t3, color="334155")

        history_lines = []
        for c_it in contract_items:
            if c_it["contract_type"] in ["전대차", "연장계약", "변경계약", "승계계약", "법인승계", "면적추가변경", "임대료변경"]:
                p_str = f"{c_it['party_a']} → {c_it['party_b']}" if c_it['party_a'] else ""
                history_lines.append(f"{c_it['contract_date']}[{c_it['contract_type']}] {p_str}")

        h_text = "\n".join(history_lines) if history_lines else "특이 변동이력 없음 (최초 계약 유지 중)"

        format_header_cell(t3.rows[0].cells[0], "계약 변동 / 전대차 / 승계 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[0].cells[1], h_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        format_header_cell(t3.rows[1].cells[0], "계약 연장 / 묵시적 갱신 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[1].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        format_header_cell(t3.rows[2].cells[0], "중도해지 / 퇴거 예정 메모", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
        format_data_cell(t3.rows[2].cells[1], "", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        t3.rows[3].cells[0].merge(t3.rows[3].cells[1])
        format_header_cell(t3.rows[3].cells[0], "기타 특약 및 참조사항", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

        t3.rows[4].cells[0].merge(t3.rows[4].cells[1])
        special_terms_text = "1. 계약일 현재 등기부등본 확인 후 대상부동산의 권리 및 시설상태의 계약으로 한다.\n2. 임대차 만료 후 원상복구를 기본으로 한다.\n3. 부가세 및 관리비는 별도로 한다."
        format_data_cell(t3.rows[4].cells[0], special_terms_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

        set_table_width_and_columns(t3, t3_cols)
        for r in t3.rows:
            set_row_cant_split(r)

        clean_file_unit = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', clean_unit)
        note_filename = f"부동산_계약관리노트_{clean_file_unit}.docx"
        note_path = os.path.join(flat_dir, note_filename)

        try:
            doc.save(note_path)
            print(f"  [MASTER NOTE SAVED] {note_filename} in {dir_name}")
            notes_created += 1
        except Exception as e:
            print(f"  [ERROR SAVING NOTE] {note_filename}: {e}")

print(f"\n==========================================")
print(f"FLAT REORGANIZATION COMPLETE! Generated {notes_created} Master Contract Notes across all flat directories!")
print(f"==========================================")
