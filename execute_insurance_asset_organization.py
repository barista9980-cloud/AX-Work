import os
import sys
import re
import shutil
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
GENERAL_AFFAIRS_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무")
INSURANCE_BASE = os.path.join(GENERAL_AFFAIRS_BASE, "03_기업보험_안전관리")
UPLOAD_DIR = os.path.join(GENERAL_AFFAIRS_BASE, r"03_보험_자산관리\00_보험_업로드_자료")

print("Executing Corporate Insurance Asset Organization Master Script...")

insurance_master_list = [
    {
        "seq": "01",
        "category": "경영인정기보험",
        "insured_object": "이종탁 대표이사",
        "company": "KDB생명",
        "policy_name": "(무)VIP경영인정기보험",
        "policy_no": "0946352030001 / 0946352030002",
        "start_date": "2022-05-31",
        "end_date": "2069-03-31",
        "period_months": 564,
        "premium": "월 10,695,000원 (건당 5,347,500원 × 2건)",
        "pay_type": "월납 (90세만기)",
        "location": "임직원 경영인 정기보장",
        "remarks": "대표이사 정기보장 보험 2건 체결",
        "pdf_match_patterns": ["KDB생명", "0946352030001", "1.폭스에듀_KDB생명"]
    },
    {
        "seq": "02",
        "category": "경영인정기보험",
        "insured_object": "이종탁 대표이사",
        "company": "매트라이프",
        "policy_name": "무배당 간편가입 Honors 경영인정기보험Plus",
        "policy_no": "13460791",
        "start_date": "2023-12-28",
        "end_date": "2069-12-28",
        "period_months": 552,
        "premium": "월 10,011,540원",
        "pay_type": "월납 (90세만기)",
        "location": "임직원 경영인 정기보장",
        "remarks": "경영인 정기보장 플랜 유지 중",
        "pdf_match_patterns": ["매트라이프", "13460791", "Honors"]
    },
    {
        "seq": "03",
        "category": "경영인정기보험",
        "insured_object": "이종탁 대표이사",
        "company": "미래에셋생명",
        "policy_name": "VIP 경영인을 위한 정기보험 무배당",
        "policy_no": "8005286685",
        "start_date": "2024-02-07",
        "end_date": "2069-02-07",
        "period_months": 540,
        "premium": "월 5,016,000원",
        "pay_type": "월납 (90세만기)",
        "location": "임직원 경영인 정기보장",
        "remarks": "미래에셋 경영인 정기보장 체결",
        "pdf_match_patterns": ["미래에셋생명", "8005286685"]
    },
    {
        "seq": "04",
        "category": "경영인정기보험",
        "insured_object": "이종탁 대표이사",
        "company": "삼성생명",
        "policy_name": "삼성 간편경영인정기보험(2403)",
        "policy_no": "41000016223329",
        "start_date": "2024-04-22",
        "end_date": "2074-04-22",
        "period_months": 600,
        "premium": "월 10,041,680원",
        "pay_type": "월납 (50년납)",
        "location": "임직원 경영인 정기보장",
        "remarks": "삼성생명 경영인 정기보장 체결",
        "pdf_match_patterns": ["삼성생명", "41000016223329", "삼성 간편경영인"]
    },
    {
        "seq": "05",
        "category": "화재재산종합보험",
        "insured_object": "대전 도룡동 204호 사업장",
        "company": "현대해상",
        "policy_name": "무배당 현대해상 성공마스터 재산종합보험",
        "policy_no": "L-025-21582789",
        "start_date": "2025-05-09",
        "end_date": "2028-05-09",
        "period_months": 36,
        "premium": "월 100,000원 (36개월 총 360만원)",
        "pay_type": "월납",
        "location": "대전 유성구 엑스포로97번길 40, 2층 204호",
        "remarks": "2022년 이전증권(L-022-04266843) 25.01.27 만기 후 25.05.09 갱신체결 완료",
        "pdf_match_patterns": ["성공마스터", "L-025-21582789", "8.주식회사 폭스에듀_현대해상", "9.주식회사 폭스에듀_현대해상"]
    },
    {
        "seq": "06",
        "category": "학원배상책임보험",
        "insured_object": "대전 스마트시티 113,114,115호 학원시설",
        "company": "DB손해보험",
        "policy_name": "학원배상책임보험",
        "policy_no": "120250591347",
        "start_date": "2025-04-30",
        "end_date": "2026-04-30",
        "period_months": 12,
        "premium": "연 20,000원 (일시납)",
        "pay_type": "일시납",
        "location": "대전 유성구 엑스포로123번길 65-38 205동 1층 113,114,115호",
        "remarks": "2024년 증권(120240729544) 만기 후 25.04.30 갱신체결 완료",
        "pdf_match_patterns": ["학원배상책임보험", "120250591347", "6.주식회사 폭스에듀_DB손해보험", "7.주식회사 폭스에듀_DB손해보험"]
    },
    {
        "seq": "07",
        "category": "홈앤비즈종합보험",
        "insured_object": "대전 스마트시티 209호 사업장",
        "company": "KB손해보험",
        "policy_name": "(무)KB홈앤비즈케어종합보험",
        "policy_no": "2024-1986936",
        "start_date": "2024-03-29",
        "end_date": "2034-03-29",
        "period_months": 120,
        "premium": "월 24,664원 (120개월 총 295만원)",
        "pay_type": "월납",
        "location": "대전 유성구 엑스포로123번길 65-38 205동 209호",
        "remarks": "10년 장기 홈앤비즈종합보험 정상 유지 중",
        "pdf_match_patterns": ["KB홈앤비즈케어", "2024-1986936", "10.주식회사 폭스에듀_KB손해보험"]
    },
    {
        "seq": "08",
        "category": "기업종합보험",
        "insured_object": "광명 GIDC A1212, A1213호 사업장",
        "company": "KB손해보험",
        "policy_name": "One KB기업종합보험",
        "policy_no": "2025-5689111",
        "start_date": "2025-08-28",
        "end_date": "2026-08-28",
        "period_months": 12,
        "premium": "연 229,000원 (일시납)",
        "pay_type": "일시납",
        "location": "경기 광명시 일직로 43, A동 1212호, 1213호",
        "remarks": "광명 GIDC 사옥 재산 및 화재 종합보장 갱신",
        "pdf_match_patterns": ["건물 화재,배상 종합보험", "2025-5689111", "11.주식회사 폭스에듀_KB손해보험"]
    },
    {
        "seq": "09",
        "category": "법인차량자동차보험",
        "insured_object": "K8 (289수3930)",
        "company": "KB손해보험 (구 현대해상)",
        "policy_name": "KB업무용 자동차보험",
        "policy_no": "2025-0127909",
        "start_date": "2025-01-08",
        "end_date": "2026-01-08",
        "period_months": 12,
        "premium": "연 1,662,600원 (일시납)",
        "pay_type": "일시납",
        "location": "K8 차량 전용 종합보험",
        "remarks": "2024년 현대해상 1,999,760원 ➔ 2025년 KB손해보험 1,662,600원 갱신",
        "pdf_match_patterns": ["K8", "289수3930", "13.K8(289수3930)"]
    },
    {
        "seq": "10",
        "category": "법인차량자동차보험",
        "insured_object": "벤츠 S클래스 (281가8991)",
        "company": "KB손해보험 (구 삼성화재)",
        "policy_name": "KB업무용 자동차보험",
        "policy_no": "2025-1512997",
        "start_date": "2025-03-10",
        "end_date": "2026-03-10",
        "period_months": 12,
        "premium": "연 1,257,880원 (일시납)",
        "pay_type": "일시납",
        "location": "벤츠 S클래스 차량 전용 종합보험",
        "remarks": "2024년 삼성화재 2,266,790원 ➔ 2025년 KB손해보험 1,257,880원 갱신",
        "pdf_match_patterns": ["벤츠", "281가8991", "15.벤츠(281가8991)"]
    },
    {
        "seq": "11",
        "category": "법인차량자동차보험",
        "insured_object": "아우디 A8 (120노2842)",
        "company": "DB손해보험",
        "policy_name": "프로미카업무용(베이직형) 자동차보험",
        "policy_no": "2-2025-6278074-000",
        "start_date": "2025-11-15",
        "end_date": "2026-11-15",
        "period_months": 12,
        "premium": "연 2,783,960원 (일시납)",
        "pay_type": "일시납",
        "location": "아우디 A8 차량 전용 종합보험",
        "remarks": "2024년 3,531,420원 ➔ 2025년 2,783,960원 DB손해보험 갱신",
        "pdf_match_patterns": ["아우디A8", "2842", "17.아우디A8(2842)"]
    },
    {
        "seq": "12",
        "category": "법인차량자동차보험",
        "insured_object": "카니발 (296더5669)",
        "company": "메리츠화재",
        "policy_name": "Readycar업무용 자동차보험",
        "policy_no": "72000-25-0259222-000",
        "start_date": "2025-03-22",
        "end_date": "2026-03-22",
        "period_months": 12,
        "premium": "연 739,470원 (배서 추납 107,420원 포함)",
        "pay_type": "일시납",
        "location": "카니발 차량 전용 종합보험",
        "remarks": "2025년 5월 23일 배서 승인 처리 완료",
        "pdf_match_patterns": ["카니발", "296더5669", "20.카니발(296더5669)"]
    }
]

# Collect all PDFs from UPLOAD_DIR
all_source_pdfs = {}
if os.path.exists(UPLOAD_DIR):
    for root, dirs, files in os.walk(UPLOAD_DIR):
        for f in files:
            if f.lower().endswith(".pdf") and not f.startswith("0.") and "계약리스트" not in f:
                fp = os.path.join(root, f)
                all_source_pdfs[f] = fp

print(f"Collected {len(all_source_pdfs)} insurance PDF documents across upload directories.")

# Clear and rebuild INSURANCE_BASE (03_기업보험_안전관리)
os.makedirs(INSURANCE_BASE, exist_ok=True)

for d in os.listdir(INSURANCE_BASE):
    dp = os.path.join(INSURANCE_BASE, d)
    if os.path.isdir(dp):
        shutil.rmtree(dp, ignore_errors=True)

print("\nBuilding 1:1 Flat Corporate Insurance Property Directory Structure in 03_기업보험_안전관리...")

for item in insurance_master_list:
    clean_obj = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', item['insured_object'])
    folder_name = f"{item['seq']}_{item['category']}_{clean_obj}({item['company']})"
    flat_folder_path = os.path.join(INSURANCE_BASE, folder_name)
    os.makedirs(flat_folder_path, exist_ok=True)

    # Match PDF
    matched_src = None
    matched_fname = None
    for pattern in item["pdf_match_patterns"]:
        for fname, fpath in all_source_pdfs.items():
            if pattern.lower() in fname.lower():
                matched_src = fpath
                matched_fname = fname
                break
        if matched_src:
            break

    raw_date = item["start_date"].replace("-", "")[2:]
    option_a_pdf = f"{item['seq']}_{item['category']}_{clean_obj}_[{item['company']}-㈜폭스에듀]_({raw_date}).pdf"

    if matched_src and os.path.exists(matched_src):
        dst_p = os.path.join(flat_folder_path, option_a_pdf)
        shutil.copy2(matched_src, dst_p)
        print(f"  [MATCHED & COPIED] {matched_fname} -> {folder_name} / {option_a_pdf}")
    else:
        print(f"  [NO PDF MATCH] {folder_name}")

# Generate Corporate Insurance Contract Notes (Word .docx)
print("\n--- GENERATING CORPORATE INSURANCE CONTRACT MANAGEMENT NOTES ---")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
    if shd is not None:
        tcPr.remove(shd)
    new_shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(new_shd)

def set_table_borders(table, color="334155", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    existing = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width_and_columns(table, col_widths_in_inches):
    total_dxa = sum([int(w * 1440) for w in col_widths_in_inches])
    tblPr = table._tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_dxa}" w:type="dxa"/>')
    existing_w = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)
    
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell_dxa = int(col_widths_in_inches[c_idx] * 1440)
            cell.width = Inches(col_widths_in_inches[c_idx])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cell_dxa}" w:type="dxa"/>')
            existing_cw = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if existing_cw is not None:
                tcPr.remove(existing_cw)
            tcPr.append(tcW)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def set_table_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def strip_keep_next(p):
    p.paragraph_format.keep_with_next = False
    pPr = p._p.get_or_add_pPr()
    kn = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}keepNext')
    if kn is not None:
        pPr.remove(kn)

def format_header_cell(cell, text, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

def format_data_cell(cell, text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF"):
    cell.text = text
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    strip_keep_next(p)
    
    if not p.runs:
        p.add_run()
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(font_size)
        r.bold = bold
        r.font.color.rgb = RGBColor(15, 23, 42)

t0_cols = [1.40, 2.025, 1.40, 2.025]
t1_cols = [0.45, 2.95, 0.85, 1.75, 0.85]
t2_cols = [1.40, 2.025, 1.40, 2.025]
t3_cols = [2.00, 4.85]

notes_created = 0

for item in insurance_master_list:
    clean_obj = re.sub(r'[\\/\:\*\?\"\<\>\|]', '_', item['insured_object'])
    folder_name = f"{item['seq']}_{item['category']}_{clean_obj}({item['company']})"
    target_folder = os.path.join(INSURANCE_BASE, folder_name)

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style_normal = doc.styles['Normal']
    style_normal.font.name = '맑은 고딕'
    style_normal.font.size = Pt(10.0)

    # P0 Title
    p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p0.text = f"[{item['category']} - {clean_obj}] 기업 보험 계약 관리 노트"
    strip_keep_next(p0)
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(2)
    for r in p0.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(13.0)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)

    # P1 Sub-note
    p1 = doc.add_paragraph()
    p1.text = "※ [Vision AI 추출 완료] 기업 보험증권 약관에서 비전 AI로 정밀 분석한 내용 입니다.\n※ 최초 작성 시 원본 PDF 확인 후 관리자 검토를 진행해 주시면 됩니다."
    strip_keep_next(p1)
    p1.paragraph_format.space_after = Pt(6)
    for r in p1.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(9.0)
        r.font.color.rgb = RGBColor(100, 116, 139)

    # Section 1 Heading
    sec1_p = doc.add_paragraph()
    sec1_p.text = "1. 보험 계약 정보"
    strip_keep_next(sec1_p)
    sec1_p.paragraph_format.space_before = Pt(6)
    sec1_p.paragraph_format.space_after = Pt(2)
    for r in sec1_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 0
    t0 = doc.add_table(rows=5, cols=4)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t0, color="334155")
    set_table_width_and_columns(t0, t0_cols)

    period_full = f"{item['start_date']} ~ {item['end_date']} ({item['period_months']}개월)"

    t0_data = [
        [("보험 종목 / 분류", True), (item["category"], False), ("피보험자 / 대상", True), (clean_obj, False)],
        [("보험 상품명", True), (item["policy_name"], False), ("보험 개시일", True), (item["start_date"], False)],
        [("보험 기간", True), (period_full, False), ("납입 방법", True), (item["pay_type"], False)],
        [("증권 번호", True), (item["policy_no"], False), ("보험료 (납입금액)", True), (item["premium"], False)],
        [("소재지 / 보장대상", True), (item["location"], False), ("계약 법인", True), ("주식회사 폭스에듀", False)]
    ]

    for r_i, r_row in enumerate(t0.rows):
        set_row_cant_split(r_row)
        row_vals = t0_data[r_i]
        for c_i, (val, is_h) in enumerate(row_vals):
            cell = r_row.cells[c_i]
            if is_h:
                format_header_cell(cell, val, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
            else:
                format_data_cell(cell, val, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    # Section 2 Heading
    sec2_p = doc.add_paragraph()
    sec2_p.text = "2. 수록 보험 증권 문서 목록 (총 1건)"
    strip_keep_next(sec2_p)
    sec2_p.paragraph_format.space_before = Pt(6)
    sec2_p.paragraph_format.space_after = Pt(2)
    for r in sec2_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 1
    t1 = doc.add_table(rows=2, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1, color="334155")
    set_table_width_and_columns(t1, t1_cols)

    set_row_cant_split(t1.rows[0])
    set_table_header_repeat(t1.rows[0])

    t1_headers = ["순서", "문서명 (파일명)", "보험 종류", "계약 당사자 (보험사 → 법인)", "보험 개시일"]
    for c_i, h in enumerate(t1_headers):
        cell = t1.rows[0].cells[c_i]
        format_header_cell(cell, h, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="E2E8F0")

    raw_d = item["start_date"].replace("-", "")[2:]
    pdf_filename_no_ext = f"{item['seq']}_{item['category']}_{clean_obj}_[{item['company']}-㈜폭스에듀]_({raw_d})"

    r_row = t1.rows[1]
    r_cells = r_row.cells
    set_row_cant_split(r_row)

    format_data_cell(r_cells[0], "01", bold=True, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    cell1 = r_cells[1]
    set_cell_background(cell1, "FFFFFF")
    set_cell_margins(cell1, top=60, bottom=60, left=100, right=100)
    cell1.text = ""
    
    p_title = cell1.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(1)
    p_title.paragraph_format.space_after = Pt(0)
    strip_keep_next(p_title)
    r_main = p_title.add_run(f"{item['policy_name']} 보험증권")
    r_main.font.name = "맑은 고딕"
    r_main.font.size = Pt(10.0)
    r_main.bold = True
    r_main.font.color.rgb = RGBColor(15, 23, 42)

    p_file = cell1.add_paragraph()
    p_file.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_file.paragraph_format.space_before = Pt(0)
    p_file.paragraph_format.space_after = Pt(1)
    strip_keep_next(p_file)
    r_sub = p_file.add_run(f"({pdf_filename_no_ext}.pdf)")
    r_sub.font.name = "맑은 고딕"
    r_sub.font.size = Pt(8.5)
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    format_data_cell(r_cells[2], item["category"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_data_cell(r_cells[3], f"{item['company']} → ㈜폭스에듀", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_data_cell(r_cells[4], item["start_date"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    set_table_width_and_columns(t1, t1_cols)

    # MANDATORY CLEAN PAGE BREAK BEFORE SECTION 3
    doc.add_page_break()

    # Section 3 Heading
    sec3_p = doc.add_paragraph()
    sec3_p.text = "3. 보험사 및 납부 정보"
    strip_keep_next(sec3_p)
    sec3_p.paragraph_format.space_before = Pt(6)
    sec3_p.paragraph_format.space_after = Pt(2)
    for r in sec3_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 2
    t2 = doc.add_table(rows=6, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2, color="334155")

    format_header_cell(t2.rows[0].cells[0], "보험회사", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[0].cells[1], item["company"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[0].cells[2], "피보험자 (법인/대표)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[0].cells[3], f"주식회사 폭스에듀 ({clean_obj})", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    format_header_cell(t2.rows[1].cells[0], "보험사 고객센터", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[1].cells[1], "대표 고객센터 / 담당FP", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[1].cells[2], "납입 형태", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[1].cells[3], item["pay_type"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    format_header_cell(t2.rows[2].cells[0], "자동이체 은행", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[2].cells[1], "법인 지정계좌", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")
    format_header_cell(t2.rows[2].cells[2], "예금주", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[2].cells[3], item["company"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    t2.rows[3].cells[0].merge(t2.rows[3].cells[1])
    t2.rows[3].cells[2].merge(t2.rows[3].cells[3])
    format_header_cell(t2.rows[3].cells[0], "보험료 납입금액", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t2.rows[3].cells[2], item["premium"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="FFFFFF")

    t2.rows[4].cells[0].merge(t2.rows[4].cells[1]).merge(t2.rows[4].cells[2]).merge(t2.rows[4].cells[3])
    format_header_cell(t2.rows[4].cells[0], "비고 (관리자 참고사항)", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

    t2.rows[5].cells[0].merge(t2.rows[5].cells[1]).merge(t2.rows[5].cells[2]).merge(t2.rows[5].cells[3])
    remarks_text = f"1. 본 건은 {item['category']} [{clean_obj}] 기업보험 관리 건임.\n2. [보장 현황] {item['remarks']}\n3. 만기 도래 1개월 전 갱신 여부 및 담보조건을 관리자가 사전 검토해 주시기 바랍니다."
    format_data_cell(t2.rows[5].cells[0], remarks_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    set_table_width_and_columns(t2, t2_cols)
    for r in t2.rows:
        set_row_cant_split(r)

    # Section 4 Heading
    sec4_p = doc.add_paragraph()
    sec4_p.text = "4. 보험 변동 및 갱신 이력"
    strip_keep_next(sec4_p)
    sec4_p.paragraph_format.space_before = Pt(6)
    sec4_p.paragraph_format.space_after = Pt(2)
    for r in sec4_p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(11.0)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 41, 59)

    # Table 3
    t3 = doc.add_table(rows=5, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3, color="334155")

    format_header_cell(t3.rows[0].cells[0], "계약 갱신 및 배서 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[0].cells[1], item["remarks"], bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    format_header_cell(t3.rows[1].cells[0], "보험금 청구 및 수령 이력", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[1].cells[1], "특이 청구이력 없음", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    format_header_cell(t3.rows[2].cells[0], "만기 / 해지 사전 준비 메모", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")
    format_data_cell(t3.rows[2].cells[1], f"만기일: {item['end_date']} (만기 1개월 전 갱신 검토)", bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    t3.rows[3].cells[0].merge(t3.rows[3].cells[1])
    format_header_cell(t3.rows[3].cells[0], "기타 특약 및 참조사항", font_size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex="F1F5F9")

    t3.rows[4].cells[0].merge(t3.rows[4].cells[1])
    special_terms_text = "1. 보험 약관상 면책사항 및 사고 통지의무를 준수함.\n2. 사업장 소재지 변경 또는 피보험자 변경 시 즉시 보험사 배서 신청 진행.\n3. 만기 도래 시 타 보험사 보장비교 후 갱신 결정."
    format_data_cell(t3.rows[4].cells[0], special_terms_text, bold=False, font_size=10.0, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex="FFFFFF")

    set_table_width_and_columns(t3, t3_cols)
    for r in t3.rows:
        set_row_cant_split(r)

    note_filename = f"기업보험_계약관리노트_{item['category']}_{clean_obj}.docx"
    note_path = os.path.join(target_folder, note_filename)

    doc.save(note_path)
    print(f"  [INSURANCE NOTE CREATED] {note_filename} in {folder_name}")
    notes_created += 1

print(f"\n==========================================")
print(f"CORPORATE INSURANCE ASSET ORGANIZATION COMPLETE! Total {notes_created} Insurance Notes Created!")
print(f"==========================================")
