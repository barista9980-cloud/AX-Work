import docx

doc_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
doc = docx.Document(doc_path)

out = []
out.append("=== PARAGRAPH RUN STYLES ===")
for i, p in enumerate(doc.paragraphs):
    for r_idx, r in enumerate(p.runs):
        font_name = r.font.name
        font_size = r.font.size.pt if r.font.size else "None"
        bold = r.bold
        color = r.font.color.rgb if r.font.color else "None"
        out.append(f"P{i} R{r_idx}: font={font_name}, size={font_size}, bold={bold}, color={color}, text='{r.text}'")

out.append("\n=== TABLE CELL DETAILS ===")
for t_idx, table in enumerate(doc.tables):
    out.append(f"\n--- TABLE {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            width = cell.width.pt if cell.width else "None"
            for p in cell.paragraphs:
                for r in p.runs:
                    f_size = r.font.size.pt if r.font.size else "None"
                    f_bold = r.bold
                    out.append(f"T{t_idx} R{r_idx} C{c_idx}: width={width}, text='{r.text}', font_size={f_size}, bold={f_bold}")

with open("template_run_details.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(out))

print("Run details exported.")
