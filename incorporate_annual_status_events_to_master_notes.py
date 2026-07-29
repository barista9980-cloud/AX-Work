import os
import sys
import re
import fitz
import docx
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"
REAL_ESTATE_BASE = os.path.join(FOXCONNECT_ROOT, r"[총무]업무\01_부동산_자산관리")
LEASE_DIR = os.path.join(REAL_ESTATE_BASE, "01_임대차계약")
SALE_DIR = os.path.join(REAL_ESTATE_BASE, "02_매매_소유권문서")

print("Processing 2022~2025 Annual Contract Status Summaries...")

# Map of extracted events from 2022~2025 PDFs
event_history = {
    "판교_판교동612": {
        "implicit": "2023-07-31 계약만기 종료 (2021-07-31 ~ 2023-07-31)",
        "termination": "",
        "remarks": "2023년 7월 31일 계약만기 종료건."
    },
    "세종_뱅크빌딩_302호": {
        "implicit": "2023-10-14 계약만기 종료",
        "termination": "2023-10-14 만기 퇴거 완료",
        "remarks": "2023년 10월 14일 만기 해지 완료."
    },
    "대전_스마트시티_2501호": {
        "implicit": "2024-11-04 묵시적갱신 (전세 12억원) / 2025-11-04 묵시적갱신 유지 중",
        "termination": "",
        "remarks": "전세 12억원 계약으로 2024, 2025년 묵시적 갱신 유지 중."
    },
    "광명_센트럴자이_1006호": {
        "implicit": "",
        "termination": "2023-03-31 중도해지 완료",
        "remarks": "2023년 3월 31일 중도해지 및 퇴거 완료."
    },
    "대전_골프존_204호,상담실": {
        "implicit": "2023-09-01 계약연장 / 2025-05-31 계약갱신 (보증금 5,000만원, 임대료 8,874,789원)",
        "termination": "2025-05-31 104호 면적 계약종료 분리 (204호+상담실 집중 유지)",
        "remarks": "2025년 5월 31일 104호 계약종료 분리 및 204호+상담실 갱신 체결."
    },
    "대전_하우스디어반_B동721호": {
        "implicit": "2023-04-07 묵시적갱신 ➔ 2024-04-07 계약연장 (1,000만/790,000원) ➔ 2025-04-07 연장계약 체결",
        "termination": "",
        "remarks": "2024년, 2025년 매년 정상 연장계약 체결 중."
    },
    "대전_케이씨씨웰츠타워_1202호": {
        "implicit": "2023-04-14 묵시적갱신 ➔ 2024-04-14 묵시적갱신 ➔ 2025-04-14 묵시적갱신 유지 중",
        "termination": "",
        "remarks": "2023년~2025년 묵시적 갱신 상태 계속 유지 중."
    },
    "광명_GIDC_1214_1215호": {
        "implicit": "2024-08-01 묵시적갱신 ➔ 2025-08-01 묵시적갱신 유지 중 (보증금 2,500만원, 월세 2,750,000원)",
        "termination": "",
        "remarks": "2024년, 2025년 묵시적 갱신 상태 유지 중."
    },
    "서초_강남역리가스퀘어_501호": {
        "implicit": "",
        "termination": "2024-11-09 만기 계약종료 (서초본점)",
        "remarks": "2024년 11월 9일 임대차 계약만기 종료."
    },
    "대전_하우스디어반_C동711호": {
        "implicit": "",
        "termination": "2024-02-19 중도해지 완료",
        "remarks": "2024년 2월 19일 중도해지 처리완료."
    },
    "대전_스타빌플러스_511호": {
        "implicit": "2024-04-08 연장계약 체결",
        "termination": "2025-04-07 만기 계약종료",
        "remarks": "2025년 4월 7일 만기 계약종료."
    }
}

print("Updating Master Contract Notes with Annual Events...")

updated_count = 0

for base_dir in [LEASE_DIR, SALE_DIR]:
    if not os.path.exists(base_dir):
        continue

    for dir_name in sorted(os.listdir(base_dir)):
        flat_dir = os.path.join(base_dir, dir_name)
        if not os.path.isdir(flat_dir) or dir_name.startswith("_"):
            continue

        docx_files = [f for f in os.listdir(flat_dir) if f.endswith(".docx")]
        if not docx_files:
            continue

        docx_path = os.path.join(flat_dir, docx_files[0])
        doc = docx.Document(docx_path)

        clean_unit = re.sub(r"^\d{2}_", "", dir_name)

        # Match event history key
        ev = None
        for k, v in event_history.items():
            if k in clean_unit or clean_unit in k:
                ev = v
                break

        if ev:
            t3 = doc.tables[3]  # Table 3 (4번 표)
            
            # Row 1: 계약 연장 / 묵시적 갱신 이력
            if ev["implicit"]:
                t3.rows[1].cells[1].text = ev["implicit"]
                p = t3.rows[1].cells[1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(10.0)

            # Row 2: 중도해지 / 퇴거 예정 메모
            if ev["termination"]:
                t3.rows[2].cells[1].text = ev["termination"]
                p = t3.rows[2].cells[1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(10.0)

            # Table 2 Row 5: 비고
            if ev["remarks"]:
                t2 = doc.tables[2]
                r_text = f"1. 본 건은 {clean_unit} 부동산 계약 관리 건임.\n2. [연도별 현황 반영] {ev['remarks']}\n3. 변경 이력 발생 시 4번 항목 표를 지속 최신화해 주시기 바랍니다."
                t2.rows[5].cells[0].text = r_text
                p = t2.rows[5].cells[0].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = "맑은 고딕"
                    r.font.size = Pt(10.0)

            doc.save(docx_path)
            print(f"  [ANNUAL EVENT REFLECTED] {dir_name} -> {docx_files[0]}")
            updated_count += 1

print(f"\n==========================================")
print(f"FINISHED REFLECTING ANNUAL ASSET EVENTS INTO {updated_count} MASTER CONTRACT NOTES!")
print(f"==========================================")
