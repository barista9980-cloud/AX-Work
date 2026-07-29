import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

template_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
doc = docx.Document(template_path)

# Set Normal Style font to 맑은 고딕 11pt
style_normal = doc.styles['Normal']
style_normal.font.name = '맑은 고딕'
style_normal.font.size = Pt(11.0)
try:
    rPr = style_normal._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="맑은 고딕" w:hAnsi="맑은 고딕" w:eastAsia="맑은 고딕" w:cs="맑은 고딕"/>')
    rPr.append(rFonts)
except Exception as e:
    print("rPr font error:", e)

# Remove ALL empty paragraphs (P4, P5, P8 etc)
for p in list(doc.paragraphs):
    if not p.text.strip():
        p._element.getparent().remove(p._element)

print("Paragraph count after purge:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    print(f"P{i}: '{p.text}'")
