import os
import docx

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

print("=== CHECKING ALL DOCX FILES IN FOXCONNECT FOR PAGE BREAK ELEMENTS ===")

count = 0
for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    for f in files:
        if f.endswith(".docx") and "계약관리노트" in f:
            full_p = os.path.join(root, f)
            try:
                doc = docx.Document(full_p)
                has_page_break = False
                for p in doc.paragraphs:
                    for r in p.runs:
                        if "\x0c" in r.text or len(r._element.xpath('.//w:br[@w:type="page"]')) > 0:
                            has_page_break = True
                            break
                    if p.paragraph_format.page_break_before:
                        has_page_break = True
                
                print(f"[{'PAGE BREAK FOUND' if has_page_break else 'SINGLE PAGE FIT'}] {f} (Tables={len(doc.tables)}, Paragraphs={len(doc.paragraphs)})")
                count += 1
            except Exception as e:
                print(f"Error checking {f}: {e}")

print(f"\nTotal notes checked: {count}")
