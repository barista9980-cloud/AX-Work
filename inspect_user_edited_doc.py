import os
import docx

fox_root = r"G:\내 드라이브\[FoxConnect]"
target_file = None

for root, dirs, files in os.walk(fox_root):
    if "402_403" in root:
        for f in files:
            if f.endswith(".docx"):
                target_file = os.path.join(root, f)
                break
    if target_file:
        break

print(f"Found target file: {target_file}")

if target_file:
    doc = docx.Document(target_file)
    out = []
    out.append("=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs):
        runs_info = []
        for r in p.runs:
            f_size = r.font.size.pt if r.font.size else "None"
            bold = r.bold
            color = r.font.color.rgb if r.font.color else "None"
            runs_info.append(f"Run('{r.text}', sz={f_size}, b={bold}, color={color})")
        out.append(f"P{i} [{p.style.name}]: text='{p.text}' | runs={runs_info}")

    out.append("\n=== TABLES ===")
    for t_idx, table in enumerate(doc.tables):
        out.append(f"\n--- TABLE {t_idx} (rows={len(table.rows)}, cols={len(table.columns)}) ---")
        for r_idx, row in enumerate(table.rows):
            row_cells = []
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                bg = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') if shd is not None else "None"
                cell_p_info = []
                for cp in cell.paragraphs:
                    align = cp.alignment
                    for cr in cp.runs:
                        f_sz = cr.font.size.pt if cr.font.size else "None"
                        bld = cr.bold
                        cell_p_info.append(f"'{cr.text}'(sz={f_sz},b={bld},align={align})")
                cell_text_clean = cell.text.replace("\n", " ").strip()
                row_cells.append(f"C{c_idx}(bg={bg}): '{cell_text_clean}' | details={cell_p_info}")
            out.append(f"  R{r_idx}: " + " | ".join(row_cells))

    with open("user_edited_402_403_details.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(out))

    print("Exported details to user_edited_402_403_details.txt")
