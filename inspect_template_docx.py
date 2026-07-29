import docx

doc_path = r"C:\Users\User\OneDrive\바탕 화면\부동산_계약관리노트_402_403호.docx"
doc = docx.Document(doc_path)

out = []
out.append("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    out.append(f"P{i} [{p.style.name}]: {p.text}")

out.append("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    out.append(f"\n--- TABLE {t_idx} (rows={len(table.rows)}, cols={len(table.columns)}) ---")
    for r_idx, row in enumerate(table.rows):
        row_text = []
        for c_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            bg = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') if shd is not None else "None"
            cell_text_clean = cell.text.replace("\n", " ").strip()
            row_text.append(f"C{c_idx}(bg={bg}): '{cell_text_clean}'")
        out.append(f"  R{r_idx}: " + " | ".join(row_text))

with open("template_structure.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(out))

print("Template structure exported to template_structure.txt")
