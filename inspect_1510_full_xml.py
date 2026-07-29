import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

FOXCONNECT_ROOT = r"G:\내 드라이브\[FoxConnect]"

file_1510_path = None
for root, dirs, files in os.walk(FOXCONNECT_ROOT):
    if "1510" in root:
        for f in files:
            if f.endswith(".docx") and "계약관리노트" in f:
                file_1510_path = os.path.join(root, f)
                break
    if file_1510_path:
        break

doc = docx.Document(file_1510_path)

for t_idx, table in enumerate(doc.tables):
    print(f"\n=======================================================")
    print(f"TABLE {t_idx} (Rows={len(table.rows)}, Cols={len(table.columns)})")
    print(f"=======================================================")
    for r_idx, row in enumerate(table.rows):
        print(f"\n--- ROW {r_idx} ---")
        for c_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            gridSpan = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
            span_val = gridSpan.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'] if gridSpan is not None else "1"
            vMerge = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
            vm_val = vMerge.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'continue') if vMerge is not None else "none"
            
            paras = [p.text for p in cell.paragraphs]
            p_str = " \n ".join(paras)
            print(f"  Cell[{c_idx}] (gridSpan={span_val}, vMerge={vm_val}): {p_str}")
